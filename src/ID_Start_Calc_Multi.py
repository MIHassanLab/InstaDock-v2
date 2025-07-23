import os
import re
import csv
import sys
import math
import shutil
import ID_utils
import tempfile
import pyKVFinder
import pandas as pd
from glob import glob
import seaborn as sns
from rdkit import Chem
from docx import Document
from rdkit.Chem import QED
from pdbfixer import PDBFixer
from datetime import datetime
from openmm.app import PDBFile
from operator import itemgetter
import matplotlib.pyplot as plt
from rdkit.Chem import Descriptors
from docx.shared import Pt, RGBColor
from cryptography.fernet import Fernet
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5.QtCore import QThread, pyqtSignal, QProcess, Qt, QEventLoop

LIGAND_TAGS = [' UNL ', ' UNK ', ' LIG ', ' <0> ', ' DRG ', ' INH ', ' NAG ', ' SO4 ', ' ATP ', ' ADP ', ' AMP ', ' HEM ', ' FMN ', ' FAD ', ' NAD ', ' GDP ', ' GTP ', ' SAM ']

STD_RESIDUES = [' ALA ', ' ARG ', ' ASN ', ' ASP ', ' CYS ', ' GLN ', ' GLU ', ' GLY ', ' HIS ', ' ILE ', ' LEU ', ' LYS ', ' MET ', ' PHE ', ' PRO ', ' SER ', ' THR ', ' TRP ', ' TYR ', ' VAL ']

STD_RESIDUES2 = ['ALA','ARG','ASN','ASP','CYS','GLN','GLU','GLY','HIS','ILE','LEU','LYS','MET','PHE','PRO','SER','THR','TRP','TYR','VAL']

key = b'XaQ5nlmxWQyPf9Q9bMjki5M7yEiXgOmTDRpvtTRWbik='

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    if hasattr(sys, '_MEIPASS'):
        # When bundled by PyInstaller, _MEIPASS contains the temporary extraction path.
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

class multiCalculateS1(QThread):
    requestResidueDialog = pyqtSignal()
    residuesReady = pyqtSignal(str)
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(list)

    def setTargetWD(self, wd):
        self.targetWD = wd

    def checkVals(self, verifylist):
        self.verifylist = verifylist

    def dockMethod(self, text):
        self.dockingMethod = text

    def findProtein(self, folder):
        files_dict = {}  # {base_name: {'pdb': path, ...}}
        
        for path in glob(os.path.join(folder, '*')):
            if not os.path.isfile(path):
                continue  # Skip directories
            
            base_name = os.path.basename(path)  # Get filename without path
            base, ext = os.path.splitext(base_name)
            ext = ext.lower()
            
            if ext in ['.pdb', '.gpf', '.pdbqt']:
                # Initialize the base entry if not exists
                if base not in files_dict:
                    files_dict[base] = {'pdb': None, 'gpf': None, 'pdbqt': None}
                
                # Perform validations (size and content checks)
                if ext in ['.pdb', '.pdbqt']:
                    proteinSize = os.path.getsize(path) / 1024
                    if proteinSize > 20:
                        with open(path, "r") as f:
                            content = f.read()
                            if any(lig in content for lig in LIGAND_TAGS) and any(res in content for res in STD_RESIDUES):
                                self.toastSignal.emit([2000, "Found Co-crystalized ligand in the protein, Continuing..."])
                
                # Assign the path to the correct file type
                files_dict[base][ext[1:]] = path  # e.g., '.pdb' → 'pdb'
                    
        return files_dict

    def run(self):
        os.chdir(self.targetWD)

        has_results = os.path.exists(os.path.join(self.targetWD, "Results"))
        has_log = os.path.exists(os.path.join(self.targetWD, "current.idlog"))
        if (has_results and not has_log):
            
            self.finishedSignal.emit([
                "❌ Results folder detected. Please remove it and try again."
            ])
            return
            
        elif has_log and not has_results:
            
            self.finishedSignal.emit([
                "❌ Error: Inconsistent state detected. Please check your files and try again."
            ])
            return
        
        elif has_log and has_results:
            
            self.toastSignal.emit([2000, "⏳ Resuming the docking process..."])
            
            self.finishedSignal.emit([
                "✅"
            ])
            return

        target_folders = list()
        main_dir_name = os.path.basename(os.path.normpath(self.targetWD)).lower()
        if main_dir_name != "backup":
            target_folders.append(self.targetWD)
        
        target_folders.extend(
            entry.path 
            for entry in os.scandir(self.targetWD) 
            if entry.is_dir() and entry.name.lower() != "backup"
        )

        valid_targets = {}

        for folder in target_folders:
            folder_files = self.findProtein(folder)

            for base, files in folder_files.items():
                # Validate based on docking method
                if self.dockingMethod == 'GPU':
                    # GPU requires at least pdb and gpf (pdbqt is optional)
                    if files['pdb'] and (files['gpf'] or files['pdbqt']):
                        valid_targets[base] = files
                    elif files['pdb'] and not files['gpf'] and not files['pdbqt']:
                        valid_targets[base] = files  # Still valid with minimal requirements
                else:
                    # Non-GPU requires at least pdb (pdbqt is optional)
                    if files['pdb']:
                        valid_targets[base] = files
                    elif (files['pdbqt'] and not files['pdb']) and self.verifylist[1]!=1:
                        valid_targets[base] = files  # Still valid with minimal requirements
                    elif (files['pdbqt'] and not files['pdb']) and self.verifylist[0] == 1:
                        error_msg = (
                            f"❌ ERROR: {base} named target protein had a PDBQT file without the corresponding PDB file. This is required for the selected site prediction method. Please verify your files and try again."
                        )
                        self.finishedSignal.emit([error_msg])
                        return

        if not valid_targets:
            
            self.finishedSignal.emit(["❌ ERROR: No valid targets found. Ensure each target has at least a PDB file."])
            return
        
        if valid_targets:
            # Handle site selection based on verifylist
            if any(x == 1 for x in self.verifylist):
                if self.verifylist[0] == 1:
                    self.requestResidueDialog.emit()
                    # 2) block *this* thread in its own QEventLoop until we get the text
                    loop = QEventLoop()
                    def _on_text(t):
                        self._residue_text = t
                        loop.quit()
                    # connect the one‐time slot
                    self.residuesReady.connect(_on_text)
                    loop.exec_()  

                    self.toastSignal.emit([2000, "Calculating bounding box dimensions\nusing user selected residues..."])
                    self.finishedSignal.emit(["Site", valid_targets, self._residue_text])
                else:
                    self.toastSignal.emit([
                        2000, 
                        "Calculating bounding box dimensions using IDv2's binding site predictor..."
                    ])
                    
                    self.finishedSignal.emit(["Predict", valid_targets, "PREDICT"])
            else:
                self.toastSignal.emit([
                    2000, 
                    "Using default blind box dimensions..."
                ])
                
                self.finishedSignal.emit(["V1", valid_targets, "BLIND"])      

                    
class multiCalculateS2(QThread):   
    finishedSignal = pyqtSignal(list)

    def completeFile(self, file):
        self.mode = file[0]
        self.proteins = file[1]
        self.tag = file[2]

    def IDv1Calc(self, proteinLocation):
        # Your existing file parsing to compute center and size.
        xval, yval, zval = [], [], []
        x, y, z = [], [], []
        with open(proteinLocation, 'r') as f:
            for line in f:
                if line.startswith('ATOM'):
                    xval.append(line[30:39])
                    yval.append(line[38:46])
                    zval.append(line[46:54])
        for i in range(len(xval)):
            x.append(float(xval[i]))
            y.append(float(yval[i]))
            z.append(float(zval[i]))

        cx = float(int((sum(x)/len(x))*1000)) / 1000.0
        lx = float(int(max(x) - min(x)) + 10)
        cy = float(int((sum(y)/len(y))*1000)) / 1000.0
        ly = float(int(max(y) - min(y)) + 10)
        cz = float(int((sum(z)/len(z))*1000)) / 1000.0
        lz = float(int(max(z) - min(z)) + 10)

        return [cx, cy, cz], [lx, ly, lz]
    
    def extract_coordinates(self, filename):
        try:
            coords = {}
            with open(filename, 'r') as file:
                for line in file:
                    # Only consider lines with an equals sign
                    if '=' in line:
                        key, value = line.split('=', 1)
                        coords[key.strip()] = value.strip()

            # Retrieve center and size values from the dictionary
            cx = coords.get('center_x')
            cy = coords.get('center_y')
            cz = coords.get('center_z')
            sx = coords.get('size_x')
            sy = coords.get('size_y')
            sz = coords.get('size_z')
            return [cx, cy, cz, sx, sy, sz]
        except:
            return ["❌ No suitable configuration file was found. Please review your files and try again."]

    def run(self):
        if self.mode == 'Site':
            proteinSites = dict()
            specificResidues = self.tag.replace(' ', '').split(',')
            resDict = ID_utils.separateMultiResults(specificResidues)
            # Get results as {"Hemoglobin": {"A": (256, 320), "B": (315, 330)}, "Cytochrome": {"B": (295, 300)}

            result = ID_utils.getMaxMinMulti(resDict)
            # Gives {'Hemoglobin': ((256, 'A'), (330, 'B')), 'Cytochrome': ((295, 'B'), (300, 'B'))}

            for target_name, files in self.proteins.items():
                pdb = files.get('pdb')
                pdbqt = files.get('pdbqt')

                for protein, mm in result.items():
                    if target_name.lower() == protein.lower():
                        minRes, maxRes = mm  # Each is a tuple: (residue, chain)
                        
                        # Choose pdbqt if available and pdb is not, otherwise use pdb
                        if pdbqt is not None and pdb is None:
                            first_coords = ID_utils.get_residue_coords(pdbqt, minRes[1], minRes[0])
                            last_coords  = ID_utils.get_residue_coords(pdbqt, maxRes[1], maxRes[0])
                        else:
                            first_coords = ID_utils.get_residue_coords(pdb, minRes[1], minRes[0])
                            last_coords  = ID_utils.get_residue_coords(pdb, maxRes[1], maxRes[0])
                        
                        center, size = ID_utils.calculate_bounding_box(first_coords, last_coords)

                proteinSites[target_name.lower()] = [files, center, size]
            
            self.finishedSignal.emit([proteinSites])

        elif self.mode == 'Predict':
            proteinSites = dict()
            pdbs = [files.get('pdb') for target_name, files in self.proteins.items()]
            for pdb in pdbs:
                pdb_path = pdb
                pdbname = os.path.basename(pdb_path).split('.')[0]
                try:
                    kv_res = pyKVFinder.run_workflow(pdb_path)
                except Exception:
                    self.finishedSignal.emit([f"❌ No cavities found by IDv2 for {pdbname}!"])
                    return
                try:
                    vols = kv_res.volume
                except Exception:
                    self.finishedSignal.emit([f"❌ No cavities found by IDv2 for {pdbname}!"])
                    return 
                cid = max(vols, key=vols.get)
                unfiltered = kv_res.residues.get(cid, [])
                if not unfiltered:
                    self.finishedSignal.emit(["❌ No residues for largest cavity."])
                    return
                
                res_list = [entry for entry in unfiltered if entry[2] in STD_RESIDUES2]
                f_res, l_res = res_list[0], res_list[-1]
                c1 = ID_utils.get_residue_coords(pdb_path, f_res[1], int(f_res[0]))
                c2 = ID_utils.get_residue_coords(pdb_path, l_res[1], int(l_res[0]))
                center, size = ID_utils.calculate_bounding_box(c1, c2)

                proteinSites[protein.split('.')[0].lower()] = [files, center, size]

            self.finishedSignal.emit([proteinSites])
            return
        else:
            proteinSites = dict()
            for target_name, files in self.proteins.items():
                pdb = files.get('pdb')
                pdbqt = files.get('pdbqt')
                # Fallback using default (blind) box dimensions.
                if pdbqt and not pdb:
                    center, size = self.IDv1Calc(pdbqt)
                else:
                    center, size = self.IDv1Calc(pdb)

                proteinSites[target_name.lower()] = [files, center, size]
            
            self.finishedSignal.emit([proteinSites])

class multiCalculateS3(QThread):
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(str)

    def setTargetWD(self, wd):
        self.targetWD = wd

    def allThings(self, text):
        self.resDict = text[0]
        #Dictionary contains target: [filesDictionary, center, size]

    def fixPDB(self, checkFix):
        self.checkFix = checkFix

    def dockMethod(self, method):
        self.method = method

    def fixPdb(self, proteinFile, targetWD):
        try:
            # Prepare a temporary filename for the fixed version.
            base_name = os.path.splitext(os.path.basename(proteinFile))[0]
            temp_fixed_file = os.path.join(os.path.dirname(proteinFile), f"{base_name}_temp_fixed.pdb")
            
            # Fix the PDB file using PDBFixer.
            fixer = PDBFixer(filename=proteinFile)
            fixer.findMissingResidues()
            fixer.findNonstandardResidues()
            fixer.replaceNonstandardResidues()
            fixer.removeHeterogens(True)
            fixer.findMissingAtoms()
            fixer.addMissingAtoms()
            
            # Write the fixed structure to a temporary file.
            with open(temp_fixed_file, 'w') as f:
                PDBFile.writeFile(fixer.topology, fixer.positions, f, keepIds=True)
            
            # Back up the original file into the main Backup directory.
            backup_dir = os.path.join(targetWD, "Backup")
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)

            # Create the backup file path and move the original there.
            backup_file = os.path.join(backup_dir, f"{base_name}_original.pdb")
            shutil.move(proteinFile, backup_file)
            
            # Replace the original file with the fixed temporary file.
            shutil.move(temp_fixed_file, proteinFile)
            
            return "Done"
        except Exception:
            return f"❌ Unable to fix the PDB, please review your files and try again."
        
    def IDv1Calc(self, proteinLocation):
        # Your existing file parsing to compute center and size.
        xval, yval, zval = [], [], []
        x, y, z = [], [], []
        with open(proteinLocation, 'r') as f:
            for line in f:
                if line.startswith('ATOM'):
                    xval.append(line[30:39])
                    yval.append(line[38:46])
                    zval.append(line[46:54])
        for i in range(len(xval)):
            x.append(float(xval[i]))
            y.append(float(yval[i]))
            z.append(float(zval[i]))

        cx = float(int((sum(x)/len(x))*1000)) / 1000.0
        lx = float(int(max(x) - min(x)) + 10)
        cy = float(int((sum(y)/len(y))*1000)) / 1000.0
        ly = float(int(max(y) - min(y)) + 10)
        cz = float(int((sum(z)/len(z))*1000)) / 1000.0
        lz = float(int(max(z) - min(z)) + 10)

        return [cx, cy, cz], [lx, ly, lz]
         
    def createConfig(self, proteinName, pdbqtFile, center, size):
        # Write the configuration file in the required format.
        with open(os.path.join(proteinName, f"{proteinName.lower()}_conf.txt"), 'w') as f:
            f.write(f"receptor = {pdbqtFile}\n\n")
            f.write(f"center_x = {center[0]:.2f}\n")
            f.write(f"center_y = {center[1]:.2f}\n")
            f.write(f"center_z = {center[2]:.2f}\n\n")
            f.write(f"size_x = {int(size[0])}\n")
            f.write(f"size_y = {int(size[1])}\n")
            f.write(f"size_z = {int(size[2])}\n")

    def find_file_in_dir(self, filename, search_dir):
        if filename is None:
            return None
        basename = os.path.basename(filename)
        pattern = os.path.join(search_dir, basename)
        return next(iter(glob(pattern)), None)
    
    def convertPDB(self, name, pdbFile, center, size):
        pdbqtFile = pdbFile.replace(".pdb", ".pdbqt")
        gpfFile = pdbFile.replace(".pdb", ".gpf")
        if self.method == "GPU":
            ar = [
                "--read_pdb", pdbqtFile,
                "-g", gpfFile, "--write_pdbqt", pdbqtFile,
                "--box_center", str(center[0]), str(center[1]), str(center[2]),
                "--box_size", str(int(size[0])), str(int(size[1])), str(int(size[2]))
            ]
        else:
            ar = [
                "--read_pdb", pdbFile,
                "-p", pdbqtFile,
                "--box_center", str(center[0]), str(center[1]), str(center[2]),
                "--box_size", str(int(size[0])), str(int(size[1])), str(int(size[2]))
            ]

        # Validate executable path
        primary_exe = resource_path("mk_prepare_receptor.exe")
        if not os.path.exists(primary_exe):
            return 0

        # Run mk_prepare_receptor
        self.process.start(primary_exe, ar)
        self.process.waitForFinished(-1)
        output = self.process.readAllStandardOutput().data().decode()
        name = (
            os.path.splitext(os.path.basename(pdbFile))[0] 
            if pdbFile 
            else os.path.splitext(os.path.basename(pdbqtFile))[0]
        )

        pdbqtName = pdbqtFile
        confPDBQT = os.path.join(self.targetWD, name, pdbqtName)
        if "error" in output.lower() or "failed" in output.lower():
            if self.method == "GPU":
                return "❌ Could not prepare the receptor molecule for the selected docking method"
            # Fallback to OpenBabel
            inputform = '-ipdb'
            alt_exe = "obabel"
            ar = [inputform, pdbFile, "-opdbqt", "-O", pdbqtName, "-xr", "-xc", "-xn", "--partialcharge", "gasteiger"]
            
            self.process.start(alt_exe, ar)
            self.process.waitForFinished(-1)
            output2 = self.process.readAllStandardOutput().data().decode()
            
            if "error" in output2.lower():
                return 0
            else:
                self.createConfig(name, confPDBQT, center, size)
                return 1
        else:
            if self.method != "GPU":
                self.createConfig(name, confPDBQT, center, size)
            return 1
        
    def run(self):
        totalCount = 0
        numTargets = len(self.resDict.keys())
        for i, (pdbName, (data_dict, center, size)) in enumerate(self.resDict.items(), start=1):
            pdb_pattern = data_dict.get("pdb")
            tPDB = next(iter(glob(pdb_pattern)), None) if pdb_pattern else None
            pdbLoc = os.path.abspath(tPDB) if tPDB is not None else None
            pdbqt_pattern = data_dict.get("pdbqt")
            tPDBQT = next(iter(glob(pdbqt_pattern)), None) if pdbqt_pattern else None
            pdbqtLoc = os.path.abspath(tPDBQT) if tPDBQT is not None else None
            gpf_pattern = data_dict.get("gpf")
            tGPF = next(iter(glob(gpf_pattern)), None) if gpf_pattern else None
            gpfLoc = os.path.abspath(tGPF) if tGPF is not None else None

            # Use targetWD as the base directory for protein_dir
            protein_dir = os.path.join(self.targetWD, pdbName)

            if pdbLoc and (not pdbqtLoc and not gpfLoc):
                with open(pdbLoc, 'r') as f:
                    lines = f.readlines()

                fixedPDB = []
                for line in lines:
                    if not line.startswith(('ATOM  ', 'HETATM')):
                        continue
                    resname = line[17:20].strip()
                    if resname == 'HOH':
                        continue
                    if resname not in STD_RESIDUES2:
                        continue
                    
                    fixedPDB.append(line)

                # Write back out
                with open(pdbLoc, 'w') as f:
                    f.writelines(fixedPDB)

                def find_altlocs(pdb_path):
                    """
                    Scan a PDB file and report any altLoc codes found.
                    """
                    altlocs = set()
                    with open(pdb_path) as f:
                        for line in f:
                            if line.startswith(("ATOM  ", "HETATM")):
                                code = line[16]              # column 17 in 0-based string indexing
                                if code not in (" ", ""):
                                    altlocs.add(code)
                    return sorted(altlocs)


                def strip_altlocs(pdb_path, output_path, keep="A"):
                    """
                    Write a new PDB with only the desired altLoc record (default “A”),
                    or all atoms that have no altLoc.
                    """
                    with open(pdb_path) as src, open(output_path, "w") as dst:
                        for line in src:
                            if line.startswith(("ATOM  ", "HETATM")):
                                code = line[16]
                                # keep if either no altLoc or exactly the one we want
                                if code in (" ", "") or code == keep:
                                    # blank out column 17 so downstream tools see no alternate locs
                                    dst.write(line[:16] + " " + line[17:])
                            else:
                                dst.write(line)

                pdb_path = pdbLoc
                altlocs = find_altlocs(pdb_path)
                if altlocs:
                    dirn, base = os.path.split(pdb_path)
                    fd, tmp_path = tempfile.mkstemp(prefix=base, dir=dirn, text=True)
                    os.close(fd)
                    strip_altlocs(pdb_path, tmp_path, keep=altlocs[0])
                    os.replace(tmp_path, pdb_path)

                self.process = QProcess()
                self.process.setProcessChannelMode(QProcess.MergedChannels)
                primary_exe = resource_path("mk_prepare_receptor.exe")
                minimal_args = ["--read_pdb", pdbLoc]
                self.process.start(primary_exe, minimal_args)
                self.process.waitForFinished(-1)
                output = self.process.readAllStandardOutput().data().decode()

                if not os.path.exists(protein_dir):
                    os.makedirs(protein_dir)
                    shutil.move(pdbLoc, protein_dir)

                pdbLoc = self.find_file_in_dir(pdbLoc, protein_dir)

                if ("error" in output.lower() or "failed" in output.lower()):
                    if self.method == "GPU":
                        self.toastSignal.emit([2000, f"Fixing and conversion in progress... ({i}/{numTargets})"])
                        testTxt = self.fixPdb(pdbLoc, self.targetWD)
                    elif self.checkFix == 1:
                        self.toastSignal.emit([2000, f"Fixing and conversion in progress... ({i}/{numTargets})"])
                        testTxt = self.fixPdb(pdbLoc, self.targetWD)
                    else:
                        self.toastSignal.emit([2000, f"Conversion in progress...({i}/{numTargets})"])
                        testTxt = "DoneWithoutFix"
                else:
                    self.toastSignal.emit([2000, f"Conversion in progress...({i}/{numTargets})"])
                    testTxt = "DoneWithoutFix"

                if testTxt == "Done":
                    name = os.path.splitext(os.path.basename(pdbLoc))[0]
                    testConv = self.convertPDB(name, pdbLoc, center, size)
                    totalCount += testConv

                elif testTxt == "DoneWithoutFix":
                    name = os.path.splitext(os.path.basename(pdbLoc))[0]
                    testConv = self.convertPDB(name, pdbLoc, center, size)
                    totalCount += testConv

            elif pdbqtLoc and (not pdbLoc and not gpfLoc):
                name = os.path.splitext(os.path.basename(pdbqtLoc))[0]
                if not os.path.exists(protein_dir):
                    os.makedirs(protein_dir)
                    shutil.move(pdbqtLoc, protein_dir)
                
                pdbqtName = name + ".pdbqt"
                confPDBQT = os.path.join(self.targetWD, name, pdbqtName)
                self.createConfig(name, confPDBQT, center, size)
                totalCount += 1

            elif pdbLoc and pdbqtLoc and not gpfLoc:
                confTest = glob(os.path.join(os.path.dirname(pdbLoc), "*conf.txt"))
                confTest = confTest[0] if confTest else None
                if not os.path.exists(protein_dir):
                    os.makedirs(protein_dir)
                    shutil.move(pdbLoc, protein_dir)
                    shutil.move(pdbqtLoc, protein_dir)
                    if confTest:
                        shutil.move(confTest, protein_dir)
                totalCount += 1
            elif pdbLoc and gpfLoc and pdbqtLoc:
                if not os.path.exists(protein_dir):
                    os.makedirs(protein_dir)
                    shutil.move(pdbLoc, protein_dir)
                    shutil.move(gpfLoc, protein_dir)
                    shutil.move(pdbqtLoc, protein_dir)
                totalCount += 1
            elif pdbLoc and gpfLoc:
                if not os.path.exists(protein_dir):
                    os.makedirs(protein_dir)
                    shutil.move(pdbLoc, protein_dir)
                    shutil.move(gpfLoc, protein_dir)
                totalCount += 1

        if totalCount == numTargets:
            self.toastSignal.emit([2000, f"Found {numTargets} targets, continuing..."])
            self.finishedSignal.emit("✅")
        else:
            self.finishedSignal.emit("❌ Failed the target conversion process, we suggest you to inspect the protein files.")


class multiCalculateS4(QThread):
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(str)
    
    def setLigandWD(self, wd):
        self.wd = wd

    def extractLibs(self, libraries):
        libmulti = None
        libsingle = None
        cwd = os.getcwd()
        ligCount = 0
        if len(libraries) > 1:
            libmulti = os.path.join(cwd, "Ligand_libraries_found")
            lib = libmulti
            if not os.path.exists(libmulti):
                os.makedirs(libmulti)
            for library in libraries:
                if library.endswith(".mol2"):
                    f = open(library, mode='r', encoding='utf-8')
                    text = f.read()
                    f.close()
                    d = "@<TRIPOS>MOLECULE"
                    splitted = [d+e for e in text.split(d) if e]
                    for i in splitted:
                        ligCount += 1
                        test = i.splitlines()
                        name = test[1]
                        outputfile =  name+".mol2"
                        u = open(outputfile,'wb')
                        file = "\n".join(test).decode('utf-8')
                        u.write(file)
                        u.close()
                elif library.endswith(".sdf"):
                    f = open(library, mode='r', encoding='utf-8')
                    text = f.read()
                    arr = text.split("$$$$")[:-1]
                    for file in arr:
                        ligCount += 1
                        if file == arr[0]:
                            name = file.splitlines()[0]
                        elif file == arr[len(arr)-1]:
                            name = file.splitlines()[1]
                        else:
                            name = file.splitlines()[1]
                        u = open(os.path.join(cwd, name+".sdf"),'wb')
                        file = "".join(file[1:len(file)-1])
                        u.write(file.encode('utf-8'))
                        u.close()
                original_library_path = os.path.join(cwd, library)
                target_path = os.path.join(libmulti, os.path.basename(library))
                shutil.move(original_library_path, target_path)

        elif len(libraries) == 1:
            library = libraries[0]
            libsingle = os.path.join(cwd, "Ligand_library_found")
            lib = libsingle
            if not os.path.exists(libsingle):
                os.makedirs(libsingle)
            if libraries[0].endswith(".mol2"):
                with open(libraries[0], mode='r', encoding='utf-8') as f:
                    text = f.read()

                d = "@<TRIPOS>MOLECULE"
                splitted = [d+e for e in text.split(d) if e]
                for i in splitted:
                    ligCount += 1
                    test = i.splitlines()
                    name = test[1]
                    outputfile =  name+".mol2"
                    with open(outputfile,'wb') as u:
                        file = "\n".join(test).decode('utf-8')
                        u.write(file)

            elif libraries[0].endswith(".sdf"):
                with open(libraries[0], mode='r', encoding='utf-8') as f:
                    text = f.read()
                arr = text.split("$$$$")[:-1]
                for file in arr:
                    ligCount += 1
                    if file == arr[0]:
                        name = file.splitlines()[0]
                    elif file == arr[len(arr)-1]:
                        name = file.splitlines()[1]
                    else:
                        name = file.splitlines()[1]
                    with open(os.path.join(cwd, name+".sdf"),'wb') as u:
                        file = "".join(file[1:len(file)-1])
                        u.write(file.encode('utf-8'))

            original_library_path = os.path.join(cwd, library)
            target_path = os.path.join(libsingle, os.path.basename(library))

            shutil.move(original_library_path, target_path)

        self.toastSignal.emit([2000, f"{ligCount} ligands extracted"])
        return lib
    
    def findLigs(self):       
        ligandFiles = list()
        libraryFiles = list()
        potentialLigands = glob('*.pdb')+ glob('*.mol2') + glob('*.sdf') + glob('*.mol') + glob('*.pdbqt')
        for ligFile in potentialLigands:
            with open(ligFile, "r") as f:
                fileContent = f.read()
            if any(res in fileContent for res in STD_RESIDUES) or "BOX" in fileContent:
                pass
            else:
                ext = os.path.splitext(ligFile)[1].lower()
                if ext == '.sdf':
                    # SDF multi‐molecule delimiter is $$$$
                    if fileContent.count('$$$$') > 1:
                        libraryFiles.append(ligFile)
                    else:
                        ligandFiles.append(ligFile)

                elif ext == '.mol2':
                    # MOL2 multi‐molecule marker is @<TRIPOS>MOLECULE
                    if fileContent.count('@<TRIPOS>MOLECULE') > 1:
                        libraryFiles.append(ligFile)
                    else:
                        ligandFiles.append(ligFile)

                else:
                    ligandFiles.append(ligFile)

        return ligandFiles, libraryFiles

    def findPDBQTLigs(self):       
        ligandFiles = list()
        
        potentialLigands = glob('*.pdbqt')
        for ligFile in potentialLigands:
            ligandFiles.append(ligFile)

        return ligandFiles

    def convertLig(self, ligFile):
        # Ensure the ligand file is an absolute path.
        ligFile = os.path.abspath(ligFile)
        
        if ligFile.endswith(".pdbqt"):
            pass
        else:       
            # Build an absolute path to the executable.
            cd = "obabel"
            format = ligFile.split('.')[1]
            inputform = "-i"+format
            tem = re.sub('\..*', '.pdbqt', ligFile)
            lig = tem
            ar = [inputform, ligFile, "-opdbqt", "-O", lig, "--partialcharge gasteiger", "-h"]
            
            self.process = QProcess()
            self.process.execute(cd, ar)
            self.process.waitForFinished(-1)

    def run(self):
        os.chdir(self.wd)
        ligandFiles, libraryFiles = self.findLigs()
        if libraryFiles:
            self.extractLibs(libraryFiles)
        ligandFiles, libraryFiles = self.findLigs()

        if ligandFiles:
            for ligandFile in ligandFiles:
                self.convertLig(ligandFile)
        else:            
            self.finishedSignal.emit("❌ No ligands found! Please check your files and try again.")
            return

        ligandFiles = self.findPDBQTLigs()

        if len(ligandFiles) > 1:
            self.toastSignal.emit([2000, f"{len(ligandFiles)} candidates ready for docking analysis!"])           
            self.finishedSignal.emit("✅")
        else:
            self.toastSignal.emit([2000, f"Ligand {ligandFiles[0]} is ready for docking analysis."])           
            self.finishedSignal.emit("✅")

class startMultiTargetDocking(QThread):
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(str)
    progressBarSignal = pyqtSignal(int)
    progressLabelSignal = pyqtSignal(str)

    def setTargetWD(self, wd):
        self.targetWD = wd
        
    def setLigandWD(self, wd):
        self.ligandWD = wd

    def checkVals(self, verifylist):
        # Set resume and logging flags based on the provided list
        self.resume = True if verifylist[0] else False
        self.logging = True if verifylist[1] else False

    def dockMethod(self, text):
        self.dockingMethod = text

    def setNConformers(self, text):
        self.conformers = text

    def resumeDocking(self, target_dir):
        log_path = os.path.join(self.targetWD, "current.idlog")
        
        if not os.path.exists(log_path):
            return []
        
        try:
            with open(log_path, "rb") as file:
                encrypted = file.read()
            decrypted = Fernet(key).decrypt(encrypted).decode("utf-8")
            log_entries = decrypted.strip().split('\n')
            
            completed_pairs = set()
            for entry in log_entries:
                try:
                    target_part, ligand_part = entry.split("::")
                    completed_pairs.add((target_part, ligand_part))
                except ValueError:
                    continue
            
            # Extract target directory base name (without extension)
            target_dir_base = os.path.basename(os.path.normpath(target_dir))
            
            completed_for_target = []
            target_dir_base = os.path.basename(os.path.normpath(target_dir))
            for target_path, ligand_part in completed_pairs:  # Split from "target_path::ligand_path"
                target_path_base = os.path.splitext(os.path.basename(target_path))[0]
                if target_path_base == target_dir_base:
                    completed_for_target.append(ligand_part)
            return completed_for_target
        
        except Exception:
            self.finishedSignal.emit(f"❌ Error parsing log")
            return []

    def process_docked_file(self, input_file, output_file):
        """
        Process an AutoDock DLG file to create a cleaned PDBQT file with modified energy remarks.
        This method removes the 'DOCKED:' prefix and replaces a 'USER' line (containing the energy)
        with a 'REMARK' line.
        """
        energy_identifier = "Estimated Free Energy of Binding"
        with open(input_file, 'r') as infile:
            lines = infile.readlines()
        
        processed_lines = []
        for line in lines:
            if line.startswith('DOCKED:'):
                clean_line = line[7:].lstrip()  # Remove 'DOCKED:' and any leading whitespace
                if clean_line.startswith('USER') and energy_identifier in clean_line:
                    # Replace 'USER' with 'REMARK' and rename the energy identifier
                    modified_line = clean_line.replace('USER', 'REMARK', 1)
                    modified_line = modified_line.replace(energy_identifier, 'Binding energy', 1)
                    processed_lines.append(modified_line)
                elif not clean_line.startswith('USER'):
                    processed_lines.append(clean_line)
        with open(output_file, 'w', newline='') as outfile:
            outfile.writelines(processed_lines)

    def findLigs(self, pdbqtFiles):
        """
        Filters the list of pdbqt files:
         - Skip files containing any standard residue (STD_RESIDUES).
         - Otherwise, consider it as a ligand file.
        """
        ligandFiles = []
        for ligFile in pdbqtFiles:
            with open(ligFile, "r") as f:
                fileContent = f.read()
                if any(res in fileContent for res in STD_RESIDUES):
                    continue
                else:
                    ligandFiles.append(ligFile)

        return ligandFiles
    
    @staticmethod
    def diff_sorted_lists(total_ligands, completed_ligands):
        """Find the last common ligand between the total and completed lists.
        Returns the sublist of `total_ligands` from the last common element onward."""
        i = j = 0
        last_common_index = -1  # Initialize to -1 to include the first element if needed

        while i < len(total_ligands) and j < len(completed_ligands):
            if total_ligands[i] == completed_ligands[j]:
                last_common_index = i  # Update to current index in total_ligands
                i += 1
                j += 1
            elif total_ligands[i] < completed_ligands[j]:
                i += 1
            else:
                j += 1

        # Return from the last_common_index (inclusive) to the end of total_ligands
        return total_ligands[last_common_index + 1:] if last_common_index != -1 else total_ligands.copy()
        
    def _save_intermediate_log(self):
        try:
            # Convert tuples to "target:ligand" strings
            log_entries = [f"{target}::{ligand}" for target, ligand in self.log]
            plain_text = "\n".join(log_entries).encode("utf-8")
            fernet = Fernet(key)
            encrypted = fernet.encrypt(plain_text)
            
            # Atomic write
            temp_path = self.temp_log_path + ".tmp"
            with open(temp_path, "wb") as temp_file:
                temp_file.write(encrypted)
            os.replace(temp_path, self.temp_log_path)
        except Exception:
            self.finishedSignal.emit(f"❌ Error saving log")
            return

    def run(self):
        self.log = []  # This will store names of ligands processed during this run
        self.temp_log_path = os.path.join(self.targetWD, "current.idlog")
        totalDockingSteps = 0 
        completed_steps = 0 
        targets_ligands = {}  # Save the list for each target

        pattern = os.path.join(self.targetWD, "**", "*.pdbqt")
        all_pdbqts = glob(pattern, recursive=True)

        targets = [
            p for p in all_pdbqts
            if all(x.lower() not in ("backup", "results") for x in p.split(os.sep))
        ]
        
        resultFolder = os.path.join(self.targetWD, "Results")
        if not os.path.exists(resultFolder):
            os.mkdir(resultFolder)

        ligFiles = [os.path.abspath(file) for file in glob(os.path.join(self.ligandWD, "*.pdbqt"))]
        total_ligands = len(ligFiles)
        total_targets = len(targets)
        sorted_total_lig_files = sorted(ligFiles, key=os.path.getsize)
        targets_ligands = {}
        completed_steps = 0
        totalDockingSteps = total_targets * total_ligands

        for target in targets:
            absTarget = os.path.abspath(target)
            baseTargetName = os.path.basename(absTarget).split(".")[0]
            baseT = os.path.basename(target).rsplit(".",1)[0]
            sorted_lig_files = sorted_total_lig_files.copy()
            
            # Apply resume filtering if needed.
            if self.resume:
                done = self.resumeDocking(os.path.join(self.targetWD, "Results", baseT))
                done_norm = {self.normalize(p) for p in done}

                full_norm = {self.normalize(p) for p in sorted_lig_files}
                if not done_norm.issubset(full_norm):
                    self.finishedSignal.emit("❌ IDLOG file compounds do not match current compounds.")
                    return

                last_idx = -1
                for i, lig in enumerate(sorted_lig_files):
                    if self.normalize(lig) in done_norm:
                        last_idx = i

                if last_idx >= 0:
                    sorted_lig_files = sorted_lig_files[last_idx:]

                completed_steps += len(done)

            targets_ligands[baseT] = sorted_lig_files

        processedDockingSteps = completed_steps

        lig_count = len(sorted_lig_files)
        #self.progressLabelSignal.emit(f"<b><center>Progress: 0% (1/{lig_count} ligand(s)) (1/{numTargets} targets)</center></b>")

        if self.dockingMethod == "AutoDock-GPU":
            autogrid_src = resource_path("autogrid4.exe")
            dlls = [resource_path("cyggcc_s-seh-1.dll"), resource_path("cyggomp-1.dll"), resource_path("cygstdc++-6.dll"), resource_path("cygwin1.dll")]
            DLLS = ["cyggcc_s-seh-1.dll", "cyggomp-1.dll", "cygstdc++-6.dll", "cygwin1.dll"]
            shutil.copy(autogrid_src, self.targetWD)
            for dll in dlls:
                shutil.copy(dll, self.targetWD)

            for target_index, target in enumerate(targets):
                targetName = os.path.basename(target).split(".")[0]
                gpfFiles = glob(os.path.join(os.path.abspath(target), "*.gpf"))
                absTarget = os.path.abspath(os.path.join(self.targetWD, targetName))
                targetName = os.path.basename(target)
                resFolder = os.path.join(resultFolder, targetName)
                if not os.path.exists(resFolder):
                    os.mkdir(resFolder)
                sorted_lig_files = targets_ligands[targetName]  # Get filtered ligands for this target
                lig_count = len(sorted_lig_files)  # Per-target ligand count

                # Check for GPF and setup
                if gpfFiles:
                    proteinFile = gpfFiles[0]
                    shutil.copy(proteinFile, resFolder)
                    protein_base = os.path.splitext(proteinFile)[0]
                    # Remove protein from ligand list (if needed)
                    sorted_lig_files = [lig for lig in sorted_lig_files if os.path.splitext(lig)[0] != protein_base]
                else:                   
                    self.finishedSignal.emit("❌ No GPF found! Please check your files and try again.")
                    return
                
                # Proceed if there are ligands to process
                if not sorted_lig_files:
                    continue
                
                # Setup grid if needed
                os.chdir(resFolder)
                proteinFile = os.path.join(resFolder, os.path.basename(gpfFiles[0]))
                proteinName = os.path.splitext(proteinFile)[0] + ".maps.fld"
                if not proteinName in os.listdir(resFolder):
                    try:
                        self.toastSignal.emit([5000, "Initializing grid parameters… this may take a while, please be patient."])
                        autogrid_exe = os.path.join(self.targetWD, "autogrid4.exe")
                        cmd = f'"{autogrid_exe}" -p "{proteinFile}"'
                        self.process = QProcess()
                        self.process.execute(cmd)
                        self.process.waitForFinished(-1)

                    except Exception:
                        self.finishedSignal.emit(f"❌ Failed to complete task.")
                        
                autodock_gpu_exe = resource_path("AutoDock-GPU.exe")
                
                self.process = QProcess()
                for j, ligand in enumerate(sorted_lig_files):
                    processedDockingSteps += 1  # Increment globally
                    rotatable_count = 0
                    with open(lig, 'r') as lig_f:
                        for line in lig_f:
                            # each REMARK line for a rotatable bond starts like "REMARK    <num>  A    between atoms:"
                            if line.startswith("REMARK") and "between atoms:" in line:
                                rotatable_count += 1

                    if rotatable_count > 32:
                        # notify & skip
                        self.toastSignal.emit([
                            2000,
                            f"Skipping {ligand}: {rotatable_count} rotatable bonds (>32) not supported on GPU."
                        ])
                        progress = (processedDockingSteps / totalDockingSteps) * 75
                        self.progressBarSignal.emit(int(progress))
                        self.progressLabelSignal.emit(
                           f"<b><center>Progress: {int(progress)}% " \
                            f"({j+1}/{lig_count} ligands) - " \
                            f"Target: {targetName}</center></b>"
                        )
                        continue
                    
                    # Docking command
                    ligName = os.path.splitext(ligand)[0]
                    finalSpot = os.path.join(resFolder, ligName)
                    cmd = f'"{autodock_gpu_exe}" --lfile "{ligand}" --ffile "{proteinName}" --nrun {self.conformers} --xmloutput 0 --N "{finalSpot}"'
                    self.process.execute(cmd)
                    self.process.waitForFinished(-1)
                    
                    # Log and save
                    self.log.append((os.path.abspath(target), os.path.abspath(ligand)))
                    self._save_intermediate_log()
                    
                    # Process DLG to PDBQT
                    self.process_docked_file(f'{finalSpot}.dlg', f'{finalSpot}_out.pdbqt')
                    
                    # Calculate progress (based on totalDockingSteps)
                    progress = (processedDockingSteps / totalDockingSteps) * 75
                    self.progressBarSignal.emit(int(progress))
                    
                    # Update label with current target's ligand progress
                    ligno = f"<b><center>Progress: {int(progress)}% " \
                            f"({j+1}/{lig_count} ligands) - " \
                            f"Target: {targetName}</center></b>"
                    self.progressLabelSignal.emit(ligno)

            os.remove(autogrid_exe)
            for DLL in DLLS:
                os.remove(os.path.join(self.targetWD, DLL))
        
        else:
            for target_index, target in enumerate(targets):
                targetName = os.path.basename(target).split(".")[0]
                resFolder = os.path.join(resultFolder, targetName)
                if not os.path.exists(resFolder):
                    os.mkdir(resFolder)
                absTarget = os.path.abspath(os.path.join(self.targetWD, targetName))
                sorted_lig_files = targets_ligands[targetName]  # Get filtered ligands for this target
                lig_count = len(sorted_lig_files)  # Per-target ligand count
                
                # Rename ligands to avoid spaces/dashes
                renamed_ligands = []
                for idx, lig in enumerate(sorted_lig_files):
                    baseLig = os.path.basename(lig)
                    new_name = baseLig.replace(" ", "_").replace("-", "")
                    if new_name != baseLig:
                        os.rename(lig, new_name)
                        renamed_ligands.append(new_name)
                    else:
                        renamed_ligands.append(lig)
                sorted_lig_files = renamed_ligands  # Use renamed list
                
                # Setup docking method and executable
                if self.dockingMethod == "AutoDock Vina":
                    exe_path = resource_path("vina.exe")
                    cueMethod = "VINA"
                elif self.dockingMethod == "QuickVina-W":
                    exe_path = resource_path("qvinaw.exe")
                    cueMethod = "QVW"
                
                self.process = QProcess()
                # Process each ligand
                for i, ligand in enumerate(sorted_lig_files):
                    processedDockingSteps += 1  # Increment globally
                    
                    # Docking command
                    outLigand = ligand.replace(".pdbqt", "_out.pdbqt")
                    confLoc = os.path.join(absTarget, f"{targetName.lower()}_conf.txt")
                    absOutLig = os.path.join(resFolder, os.path.basename(outLigand))
                    
                    # Build command with logging (if applicable)
                    cmd = f'"{exe_path}" --config "{confLoc}" --ligand "{ligand}" --num_modes {self.conformers} --out "{absOutLig}"'
                    if self.logging and self.dockingMethod == "QuickVina-W":
                        loglig = re.sub(r'\.pdbqt$', '_log.log', ligand)
                        absLogLig = os.path.join(resFolder, os.path.basename(loglig))
                        cmd = f'"{exe_path}" --config "{confLoc}" --ligand "{ligand}" --log "{absLogLig}" --num_modes {self.conformers} --out "{absOutLig}"'
                    
                    # Execute and log
                    self.process.execute(cmd)
                    self.process.waitForFinished(-1)
                    self.log.append((os.path.abspath(target), os.path.abspath(ligand)))
                    self._save_intermediate_log()
                    
                    # Calculate progress
                    progress = (processedDockingSteps / totalDockingSteps) * 75
                    self.progressBarSignal.emit(int(progress))
                    
                    # Update label with current target's ligand progress
                    ligno = f"<b><center>Progress: {int(progress)}% " \
                            f"({i+1}/{lig_count} ligands) - " \
                            f"Target: {targetName}</center></b>"
                    self.progressLabelSignal.emit(ligno)

        cue = f'DONE-{cueMethod}'
        
        self.finishedSignal.emit(cue)

        if self.log:
            current_time = datetime.now()
            formatted_time = current_time.strftime("%Y-%m-%d_%H-%M-%S")
            final_log_path = os.path.join(self.targetWD, f"{formatted_time}.idlog")
            try:
                os.rename(self.temp_log_path, final_log_path)
            except FileNotFoundError:
                # Fallback if temp file is missing (first run)
                with open(final_log_path, "wb") as file:
                    fernet = Fernet(key)
                    file.write(fernet.encrypt('\n'.join([f"{t}:{l}" for t,l in self.log]).encode()))

class continueMultiID(QThread): 
    progressBarSignal = pyqtSignal(int)
    progressLabelSignal = pyqtSignal(str)
    finishedSignal = pyqtSignal(list)
    toastSignal = pyqtSignal(list)

    def setTargetWD(self, wd):
        # Set self.wd as the parent folder where the Results folder (and its subfolders) will be created.
        self.targetWD = wd

    def setLigandWD(self, wd):
        # Set self.wd as the parent folder where the Results folder (and its subfolders) will be created.
        self.ligandWD = wd

    def setDPI(self, dpi):
        self.dpi = int(dpi)

    def setBlind(self, blind):
        self.blind = blind

    def setBinding(self, binding):
        self.binding = binding

    def setSiteSpecific(self, siteSpecific):
        self.siteSpecific = siteSpecific

    def dockMethod(self, method):
        # Set the method description along with citation numbers.
        if method == "DONE-VINA":
            self.method = "AutoDock Vina [2]"
        elif method == "DONE-QVW":
            self.method = "QuickVina-W [2] (modified AutoDock Vina [3])"
        elif method == "DONE-GPU":
            self.method = "AutoDock-GPU [2]"

    def setLogging(self, log):
        self.logging = log

    def findTarget(self):
        if self.method == "AutoDock-GPU [2]":
            proteinName = glob("*.maps.fld")
            return proteinName
        else:
            with open("conf.txt", 'r') as f:
                for line in f:
                    if line.startswith("receptor"):
                        return line.split("=")[1].strip()
                    
    def moveTarget(self, receptorName):
        target_dir = self.wd + f"/{receptorName} related files"
        if not os.path.exists(str(target_dir)):
            os.mkdir(str(target_dir))
        for pattern in [f"{receptorName}*", "boron-silicon-atom_par.dat"]:
            for file in glob(pattern):
                if file not in f"{receptorName} related files":
                    shutil.move(file, str(target_dir) + os.sep + os.path.basename(file))        
                    
    def human_format(num):
        suffixes = [
            (1e12, 'Trillion'),
            (1e9, 'Billion'),
            (1e6, 'Million'),
            (1e3, 'Thousand'),
        ]
        for val, suffix in suffixes:
            if num >= val:
                formatted_num = num / val
                # Use .3g to format up to 3 significant figures
                return f"{formatted_num:.3g} {suffix}"
        return str(num)

    def create_docking_document(self, allInfo, results_dir, mode, dockMethod):
        #receptor_name = full with format
        #results_dir = complete results path
        #mode = vs or docking
        #top_hits = selected top hits
        #dockMethod = "AutoDock Vina [2]", "QuickVina-W [2] (modified AutoDock Vina [3])" OR "AutoDock-GPU [2]"
        #Ligands contains a list of dictionaries with all ligand information.
        # Example.:'Name': LigandName, 'Binding Energy (kcal/mol)': -6.0, 'pKi': 4.401, 'Ligand Efficiency (kcal/mol/non-H atom)': 0.051, 'Torsional Energy': 1.245, 'Molecular Weight': 16.0, 'HBD': 0, 'HBA': 0, 'LogP': 0.6, 'Lipinski Violations': '0', 'Lipinski Compliant': 'True'
    
        all_dicts = [entry for sublist in allInfo for entry in sublist]

        keys_of_interest = [
            "Target",
            "Ligand",
            "Binding Free Energy (kcal/mol)",
            "pKi",
            "Ligand Efficiency (kcal/mol/non-H atom)",
            "Torsional Energy (kcal/mol)",
        ]

        unique_values = {}
        for key in keys_of_interest:
            # Collect all values for the key and deduplicate
            values = [d.get(key, "N/A") for d in all_dicts]  # Use .get() for safety
            unique_values[key] = list(dict.fromkeys(values))

        document = Document()
        styles = document.styles
        receptor_names = unique_values["Target"]
        ligand_names = unique_values["Ligand"]
        binding_energies = unique_values["Binding Free Energy (kcal/mol)"]
        pkis = unique_values["pKi"]
        ligand_efficiencies =unique_values["Ligand Efficiency (kcal/mol/non-H atom)"]
        torsional_energies = unique_values["Torsional Energy (kcal/mol)"]

        records = list(zip(receptor_names, ligand_names, binding_energies, pkis, ligand_efficiencies, torsional_energies))

        records_sorted = sorted(records, key=lambda x: x[2])

        lowest_record = records_sorted[0]   # Most negative binding energy
        highest_record = records_sorted[-1]   # Binding energy closest to zero

        style_one = styles.add_style('Heading One', WD_STYLE_TYPE.PARAGRAPH)
        style_one.base_style = styles['Heading 1']
        font = style_one.font
        font.name = 'Times New Roman'
        font.size = Pt(16)
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)

        style_three = styles.add_style('Heading Three', WD_STYLE_TYPE.PARAGRAPH)
        style_three.base_style = styles['Heading 2']
        font = style_three.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = False
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)     

        number_of_receptor = len(receptor_names)
        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            methodDescription = "a hybrid scoring function (empirical + knowledge-based) in docking calculations"
        elif dockMethod == "AutoDock Vina [2]":
            methodDescription = "a scoring function based on the Vina algorithm"
        elif dockMethod == "AutoDock-GPU [2]":
            methodDescription = "a scoring function based on the AutoDock4 algorithm"

        if self.blind and not self.binding:
            approach = "a blind search space for the ligand"
        if self.blind and self.binding:
            approach = "a predicted search space for the ligand"
        if self.siteSpecific:
            approach = "a site-specific search space for the ligand"

        modeToWrite = "-based virtual screening"
        
        if len(ligand_names) == 1:
            ligand = ligand_names[0]
            para = document.add_paragraph(f'Molecular docking of \"{ligand}\" with {number_of_receptor} target molecules', style = 'Heading One')
            pformat = para.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = document.add_paragraph(f'Molecular docking of \"{ligand}\" with {number_of_receptor} target molecules was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
            p.add_run(f"The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}")
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif len(ligand_names) == 2:
            ligand1 = ligand_names[0]
            ligand2 = ligand_names[1]
            para = document.add_paragraph(f'Molecular docking of \"{ligand1}\" and \"{ligand2}\" with {number_of_receptor} target molecules', style = 'Heading One')
            pformat = para.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = document.add_paragraph(f'Molecular docking of \"{ligand1}\" and \"{ligand2}\" with {number_of_receptor} target molecules was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
            p.add_run(f'The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}s.')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif len(ligand_names) > 2:
            ligLen = len(ligand_names)
            if ligLen > 10000:
                ligLen = self.human_format(len(ligand_names))
            para = document.add_paragraph(f'Virtual screening of {ligLen} compounds with {number_of_receptor} target molecules', style = 'Heading One')
            pformat = para.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = document.add_paragraph(f'Virtual screening of {ligLen} compounds with {number_of_receptor} target molecules was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
            p.add_run(f'The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}s.')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        style_two = styles.add_style('Heading Two', WD_STYLE_TYPE.PARAGRAPH)
        style_two.base_style = styles['Heading 2']
        font = style_two.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)
        paragraph = document.add_paragraph('Materials and Methods', style = 'Heading Two')
        pformat = paragraph.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        p = document.add_paragraph('The ', style='Heading Three')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.add_run('p')
        p.add_run('Ki').italic = True
        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            p.add_run(', the negative decimal logarithm of inhibition constant [4] was calculated from the ∆')
        else:
            p.add_run(', the negative decimal logarithm of inhibition constant [3] was calculated from the ∆')
        p.add_run('G').italic = True
        p.add_run(' parameter while using the following formula:')
        style_six = styles.add_style('Heading Six', WD_STYLE_TYPE.PARAGRAPH)
        style_six.base_style = styles['Heading 6']
        font = style_six.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        font.bold = False
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)
        d = document.add_paragraph('', style='Heading Three')
        dformat = d.paragraph_format
        dformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        d.add_run('∆')
        d.add_run('G').italic = True
        d.add_run(' = RT(Ln ')
        d.add_run('Ki').italic = True
        test = d.add_run('pred')
        test.font.subscript = True
        d.add_run(')')
        run = d.add_run()
        run.add_break()
        d.add_run('Ki').italic = True
        test = d.add_run('pred')
        test.font.subscript = True
        d.add_run(' = e')
        test = d.add_run('(∆')
        test.font.superscript = True
        test = d.add_run('G')
        test.font.superscript = True
        test.italics = True
        test = d.add_run('/RT)')
        test.font.superscript = True
        run = d.add_run()
        run.add_break()
        d.add_run('p')
        d.add_run('Ki').italic = True
        d.add_run(' = -log(')
        d.add_run('Ki').italic = True
        test = d.add_run('pred')
        test.font.subscript = True
        d.add_run(')')
        p = document.add_paragraph('where ∆', style='Heading Three')
        p.add_run('G').italic = True
        p.add_run(' is the binding affinity (kcal mol')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run('), R (gas constant) is 1.98 cal*(mol*K)')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run(', T (room temperature) is 298.15 Kelvin, and ')
        p.add_run('Ki').italic = True
        test = p.add_run('pred')
        test.font.subscript = True
        p.add_run(' is the predicted inhibitory constant.')

        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            p = document.add_paragraph("Ligand efficiency (LE) is a commonly applied parameter for selecting favorable ligands by comparing the values of average binding energy per atom [5]. The following formula was applied to calculate LE:", style='Heading Three')
        else:
            p = document.add_paragraph("Ligand efficiency (LE) is a commonly applied parameter for selecting favorable ligands by comparing the values of average binding energy per atom [4]. The following formula was applied to calculate LE:", style='Heading Three')
        p.add_run("")
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        d = document.add_paragraph('', style='Heading Three')
        dformat = d.paragraph_format
        dformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        d.add_run('LE = -∆')
        d.add_run('G').italic = True
        d.add_run('/N')
        p = document.add_paragraph('where LE is the ligand efficiency (kcal mol', style='Heading Three')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run(' non-H atom')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run('), ∆')
        p.add_run('G').italic = True
        p.add_run(' is binding affinity (kcal mol')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run(') and N is the number of non-hydrogen atoms in the ligand.')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            p = document.add_paragraph(
                "Quantitative Estimate of Drug-likeness (QED) assesses how 'drug-like' a molecule is "
                "by combining multiple physicochemical properties [6]. The formula being:",
                style='Heading Three'
            )
        else:
            p = document.add_paragraph(
                "Quantitative Estimate of Drug-likeness (QED) assesses how 'drug-like' a molecule is "
                "by combining multiple physicochemical properties [5]. The formula being:",
                style='Heading Three'
            )
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # render the QED equation
        qed_eq = document.add_paragraph('', style='Heading Three')
        qed_eq_format = qed_eq.paragraph_format
        qed_eq_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qed_eq_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Construct QED = (d₁ʷ¹ × d₂ʷ² × ... × d₈ʷ⁸)^(1 / ∑wᵢ)
        qed_eq.add_run('QED = (')
        qed_eq.add_run('d')
        qed_eq.add_run('1').font.subscript = True
        qed_eq.add_run('w1').font.superscript = True
        qed_eq.add_run(' x ')
        qed_eq.add_run('d')
        qed_eq.add_run('2').font.subscript = True
        qed_eq.add_run('w2').font.superscript = True
        qed_eq.add_run(' x ')
        qed_eq.add_run('...')
        qed_eq.add_run(' x ')
        qed_eq.add_run('d')
        qed_eq.add_run('8').font.subscript = True
        qed_eq.add_run('w8').font.superscript = True
        qed_eq.add_run(')')
        qed_eq.add_run('1/∑').font.superscript = True
        p = qed_eq.add_run('wi')
        p.font.italic = True
        p.font.superscript = True

        # a brief “where …” description
        p = document.add_paragraph(
            "where each dᵢ is the desirability of property i, and wᵢ is its corresponding weight.",
            style='Heading Three'
        )
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        p = document.add_paragraph('Results', style = 'Heading Two')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if len(ligand_names) == 1:
            ligand = ligand_names[0]
            bEnergy = lowest_record[2]
            pKi = lowest_record[3]
            lEfficiency = lowest_record[4]
            p = document.add_paragraph(f'Compound \"{ligand}\" was subjected to docking analysis, it presented the best predicted binding free energy of {bEnergy} kcal mol', style='Heading Three')
            test = p.add_run('-1')
            test.font.superscript = True
            p.add_run(' and a p')
            p.add_run('Ki').italic = True
            p.add_run(f' value of {pKi} towards the receptor {lowest_record[0]}. It also possesses a ligand efficiency of {lEfficiency} kcal mol ')
            test = p.add_run('-1')
            test.font.superscript = True
            p.add_run(' non-H atom')
            test = p.add_run('-1')
            test.font.superscript = True
            p.add_run(f'. Please refer to the CSV file given in the central directory for other docking parameters of the compound used in this study.')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif len(ligand_names) == 2:
            ligand1 = ligand_names[0]
            ligand2 = ligand_names[1]
            bEnergy = lowest_record[2]
            pKi = lowest_record[3]
            lEfficiency = lowest_record[4]
            
            p = document.add_paragraph(f'Compound \"{ligand1}\" and \"{ligand2}\" were subjected to docking analysis. Compound \"{lowest_record[1]}\" presented the better predicted binding free energy of {bEnergy} kcal/mol towards the receptor {lowest_record[0]}. Please refer to the CSV file given in the central directory for other docking parameters of each compound used in this study.', style='Heading Three')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        else:
            p = document.add_paragraph(f'All {ligLen} compounds were subjected to docking analysis and presented a predicted binding free energy within the range of {lowest_record[2]} kcal/mol towards {lowest_record[0]} target molecule, and {highest_record[2]} kcal/mol towards {highest_record[0]}. Please refer to the CSV file given in the central directory for the binding affinities and other docking parameters of each compound used in this study, where the best predicted binding free energy was observed in the case of {lowest_record[1]} as {lowest_record[2]} kcal/mol.', style='Heading Three')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        p = document.add_paragraph('References', style = 'Heading Two')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        style_four = styles.add_style('Heading Four', WD_STYLE_TYPE.PARAGRAPH)
        style_four.base_style = styles['List Number']
        font = style_four.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = False
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)

        c = document.add_paragraph('Mathur Y, Hassan MI (2025). InstaDock-v2: TBD. ', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run(' TBD').italic = True #Journal
        c.add_run(', ')
        c.add_run('TBD').bold = True #Volume
        c.add_run('TBD') #Pages and DOI
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            c = document.add_paragraph('Hassan NM, Alhossary AA, Mu Y, Kwoh CK (2017). Protein-ligand blind docking using QuickVina-W with inter-process spatio-temporal integration,' , style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Scientific Reports').italic = True
            c.add_run(', ')
            c.add_run('7(1)').bold = True
            c.add_run(': 1-13.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            c = document.add_paragraph('Jerome Eberhardt, Diogo Santos-Martins, Andreas F. Tillack, and Stefano Forli (2021) AutoDock Vina 1.2.0: New Docking Methods, Expanded Force Field, and Python Bindings.', style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Journal of Chemical Information and Modeling').italic = True
            c.add_run(', ')
            c.add_run('61(8)').bold = True
            c.add_run(': 3891-3898.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif dockMethod == "AutoDock Vina [2]":
            c = document.add_paragraph('Jerome Eberhardt, Diogo Santos-Martins, Andreas F. Tillack, and Stefano Forli (2021) AutoDock Vina 1.2.0: New Docking Methods, Expanded Force Field, and Python Bindings.', style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Journal of Chemical Information and Modeling').italic = True
            c.add_run(', ')
            c.add_run('61(8)').bold = True
            c.add_run(': 3891-3898.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif dockMethod == "AutoDock-GPU [2]":
            c = document.add_paragraph('Diogo Santos-Martins, Leonardo Solis-Vasquez, Andreas F Tillack, Michel F Sanner, Andreas Koch, and Stefano Forli (2021) Accelerating AutoDock4 with GPUs and Gradient-Based Local Search.', style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Journal of Chemical Theory and Computation').italic = True
            c.add_run(', ')
            c.add_run('61(8)').bold = True
            c.add_run(': 3891-3898.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 

        c = document.add_paragraph('Shityakov, S., & Förster, C. (2014). In silico structure-based screening of versatile P-glycoprotein inhibitors using polynomial empirical scoring functions.', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run(' Advances and Applications in Bioinformatics and Chemistry').italic = True
        c.add_run(', ')
        c.add_run('7: ').bold = True
        c.add_run('1.')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('Hopkins, A. L., Groom, C. R., & Alex, A. (2004). Ligand efficiency: a useful metric for lead selection.', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run('  Drug discovery today').italic = True
        c.add_run(', ')
        c.add_run('9(10)').bold = True
        c.add_run(': 430.')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('Bickerton, G. R., Paolini, G. V., Besnard, J., Muresan, S., & Hopkins, A. L. (2012). Quantifying the chemical beauty of drugs.', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run('  Nature chemistry').italic = True
        c.add_run(', ')
        c.add_run('4(2)').bold = True
        c.add_run(': 90-98.')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('______________________________')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

        c = document.add_paragraph('', style = 'Heading Three')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        j = c.add_run('Attention user: ')
        j.bold = True
        j.italic = True
        c.add_run(' This is an auto-generated writeup for your job, please crosscheck this with the generated output.').italic = True
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        style_five = styles.add_style('Heading Five', WD_STYLE_TYPE.PARAGRAPH)
        style_five.base_style = styles['Heading Five']
        font = style_five.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = True
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)
        paragraph = document.add_paragraph('-Team IDv2',style = 'Heading Five')
        pformat = paragraph.paragraph_format
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pformat.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.save(os.path.join(results_dir, 'IDv2 Result Summary.docx'))
        return document
    
    def pdbqt_to_qed(self, pdbqt_path):
        try:
            proc = QProcess()
            args = ["-ipdbqt", pdbqt_path, "-osmi", "-xn"]
            cd = "obabel"
            proc.start(cd, args)
            proc.waitForFinished(-1)

            stdout = proc.readAllStandardOutput().data().decode().strip()
            stderr = proc.readAllStandardError().data().decode().strip()

            # obabel outputs lines like: "SMILES  name"
            lines = stdout.strip().splitlines()

            # take only the very first line
            first_line = lines[0]
            smiles = first_line.split()[0]

            # Build RDKit Mol and compute QED
            mol = Chem.MolFromSmiles(smiles, sanitize=True)
            mol = Chem.AddHs(mol)
            try:
                Chem.SanitizeMol(mol)
            except Chem.SanitizeException:
                pass    
            resultQED = QED.qed(mol)

        except Exception as a:
            resultQED = 0.0

        return resultQED

    def create_affinity_file_from_pdbqt(self, input_dir, pdbqt_files):
        results = []
        for pdbqt_file in pdbqt_files:
            # Initialize entry with ligand name
            ligand_entry = {'Target': os.path.basename(input_dir)}
            ligand_entry.update({'Ligand': os.path.splitext(os.path.basename(pdbqt_file))[0].replace('_out', '')})

            # Read PDBQT content
            with open(pdbqt_file, 'r') as f:
                content = f.readlines()

            # Extract torsional energy
            energy_value = None
            # Compute torsional energy using the FIRST occurrence of TORSDOF
            torsional_energy = 'NA'
            try:
                tors_line = next(line for line in content if "TORSDOF" in line)
                # Assume the line format is: "TORSDOF   <number>"
                tors_value = float(tors_line.split()[1])
                torsional_energy = round(tors_value * 0.3113, 3)
            except (StopIteration, IndexError, ValueError):
                torsional_energy = 'NA'

            try:
                energy_line = next(line for line in content 
                                   if line.startswith('REMARK') and 'Binding energy' in line)
                energy_value = float(energy_line.split('=')[1].strip().split()[0])
            except (StopIteration, IndexError, ValueError):
                try:
                    energy_line = next(line for line in content if line.startswith('REMARK VINA RESULT'))
                    parts = energy_line.strip().split()
                    energy_value = float(parts[3])
                except (StopIteration, IndexError, ValueError):
                    self.toastSignal.emit([2000, f"Skipping {pdbqt_file}: No valid energy line found"])
                    continue

            # Count heavy (non-H) atoms
            nh_count = sum(1 for line in content if line.startswith(('ATOM', 'HETATM')) and ' H ' not in line)
            ligand_entry.update({
                'Binding Free Energy (kcal/mol)': energy_value,
                'Ligand Efficiency (kcal/mol/non-H atom)': round(-energy_value / nh_count, 3) if nh_count else 0,
                'pKi': round(-math.log10(math.exp(energy_value * 1000 / (1.986 * 298.15))), 3),
                'Torsional Energy (kcal/mol)': torsional_energy
            })

            ligand_entry.update({'QED': self.pdbqt_to_qed(pdbqt_file)})
            results.append(ligand_entry)

        return results

    def findLigs(self, location):
        """Return list of ligand files matching the pattern *_out.pdbqt."""
        return glob(os.path.join(location, '*_out.pdbqt'))

    def move_ligand_files(self):
        """
        Moves all *_out.pdbqt and (if self.logging is True) *_log.log files from the current working directory
        to a folder structure under the Results folder.
        Returns the path to the Docked_ligands folder.
        """
        results_dir = self.wd + "/Results"
        docked_dir = results_dir + "/Docked Ligands"
        if not os.path.exists(str(docked_dir)):
            os.mkdir(str(docked_dir))
        if self.logging:
            docking_log_dir = docked_dir + "/Docking Log"
            if not os.path.exists(str(docking_log_dir)):
                os.mkdir(str(docking_log_dir))
        for pattern in ["*_out.pdbqt", "*_log.log"]:
            for file in glob(pattern):
                if file.endswith("_log.log"):
                    if self.logging:
                        shutil.move(file, str(docking_log_dir) + os.sep + os.path.basename(file))
                elif file.endswith("_out.pdbqt"):
                    shutil.move(file, str(docked_dir) + os.sep + os.path.basename(file))
        return docked_dir

    def create_top_hits_folder(self, ligand_info_list, receptor_name, top_hits):
        """
        Creates the top hits folder in the Results/Top hits(With Poses) folder by selecting the best compounds,
        copying their pdbqt files (from Docked_ligands) into that folder, splitting each pdbqt file into individual
        conformer files (if multiple conformers exist), and creating an affinity CSV file only for the main pdbqt file.
        Returns a tuple containing the path to the top hits folder and the top hits affinity CSV file.
        """
        results_dir = self.wd + "/Results"
        split_dir = results_dir + "/Top hits (With Poses)"
        if not os.path.exists(str(split_dir)):
            os.mkdir(str(split_dir))
        
        sorted_ligs = sorted(ligand_info_list, key=itemgetter('Binding Free Energy (kcal/mol)'))
        top_ligs = sorted_ligs[:top_hits]
        docked_dir = results_dir + "/Docked Ligands"
        
        for lig in top_ligs:
            src = docked_dir + f"/{lig['Name']}_out.pdbqt"
            if os.path.exists(src):
                dst = str(split_dir) + os.sep + os.path.basename(src)
                shutil.copy(src, dst)
                with open(dst, 'r') as f:
                    content = f.read()
                models = [m.strip() for m in content.split("ENDMDL") if m.strip() != ""]
                if len(models) > 1:
                    base = os.path.splitext(os.path.basename(src))[0]
                    os.remove(dst)
                    for i, model in enumerate(models, 1):
                        conf_file = os.path.join(str(split_dir), f"{base}_ligand_{i}.pdbqt")
                        with open(conf_file, 'w') as cf:
                            cf.write(model + "\nENDMDL\n")
                    main_file = os.path.join(str(split_dir), f"{base}_ligand_1.pdbqt")
                else:
                    main_file = dst
                lig['main_file'] = main_file

        top_csv_path = str(split_dir) + os.sep + "IDv2 Results (Top hits).csv"
        with open(top_csv_path, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Name of the ligand", "Binding Free Energy (kcal/mol)", "pKi",
                             "Ligand Efficiency (kcal/mol/non-H atom)", "Torsional Energy (kcal/mol)", "QED"])
            for lig in top_ligs:
                main_file = lig.get('main_file', "")
                if not main_file or not os.path.exists(main_file):
                    continue
                with open(main_file, 'r') as j:
                    ha = j.read()
                ha_main = ha.split("ENDMDL", 1)[0]
                ha_lines = ha_main.splitlines()
                try:
                    energy_line = ha_lines[1]
                    ter = float(energy_line.split(':')[1].split()[0])
                    self.gridLayout1.addWidget(self.browseWDButton, 0, 1, 1, 1, Qt.AlignRight)

                except Exception:
                    ter = lig.get("Binding Free Energy (kcal/mol)", 0)
                t = lig.get("Torsional Energy (kcal/mol)", "NA")
                ki = math.exp(ter * 1000 / (1.986 * 298.15))
                pki = -math.log(ki, 10)
                pki = round(pki, 2)
                nhcount = 0
                for line in ha_lines:
                    if line.startswith('ATOM') or line.startswith('HETATM'):
                        if ' H ' not in line:
                            nhcount += 1
                    else:
                        nhcount -= 1
                le = -(ter / nhcount) if nhcount != 0 else 0
                le = round(le, 4)
                ligand_name = re.sub(r'\_out\.pdbqt$', '', os.path.basename(main_file))
                writer.writerow([ligand_name, ter, pki, le, t])
        return split_dir, top_csv_path

    def create_binding_energy_plot(self, affinity_vals, output_path, dpi):
        """
        Create a histogram with KDE showing the distribution of binding energies.
        """
        # Convert to DataFrame with proper column names
        df = pd.DataFrame(affinity_vals, columns=["Target", "Binding Free Energy (kcal/mol)"])

        plt.figure(figsize=(10, 6), dpi=dpi)
        sns.histplot(
            data=df,
            x="Binding Free Energy (kcal/mol)",
            hue="Target",
            multiple="stack",
            kde=True,
            bins=15,
            palette="viridis"
        )
        
        plt.title("Binding Energy Distribution by Target")
        plt.xlabel("Binding Free Energy (kcal/mol)")
        plt.ylabel("Frequency")
        plt.tight_layout()
        
        # Save and close
        plt.savefig(output_path, dpi=dpi, bbox_inches="tight")
        plt.close()

    def run(self):
        """
        Main method executed in this thread.
        Steps:
          1. Find docking result ligand files.
          2. Create the overall affinity CSV file (saved in Results folder).
          3. Create the docking document in Results.
          4. Create plot in Results.
        """
        results_dir = os.path.join(self.targetWD, "Results") 
        os.chdir(results_dir)
        subfolders = [entry.name for entry in os.scandir() if entry.is_dir()]
        
        affinity_vals = list()
        self.progressBarSignal.emit(75)
        self.progressLabelSignal.emit("<b><center>Progress: 75% (Creating overall affinity results file...)</center></b>")
        fieldnames = ['Target', 'Ligand', 'Binding Free Energy (kcal/mol)', 'pKi', 'Ligand Efficiency (kcal/mol/non-H atom)', 'Torsional Energy (kcal/mol)', 'QED']
        for target in subfolders:
            pdbqt_files = self.findLigs(os.path.join(results_dir, target))
            if not pdbqt_files:
                self.finishedSignal.emit(["❌ No docking result files found."])
                return
            affinity_vals.append(self.create_affinity_file_from_pdbqt(target, pdbqt_files))

        self.progressBarSignal.emit(80)
        self.progressLabelSignal.emit("<b><center>Progress: 80% (Parsing overall affinity results...)</center></b>")
        with open(os.path.join(results_dir, 'Combined_IDv2_Results.csv'), 'w', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for affinity_val in affinity_vals:
                for row in affinity_val:
                    # Process each value: if it's a float, round it; otherwise, leave it as is.
                    processed_row = {k: round(v, 3) if isinstance(v, float) else v for k, v in row.items()}
                    writer.writerow(processed_row)

        self.mode = "virtual_screening"
        self.progressBarSignal.emit(85)
        self.progressLabelSignal.emit("<b><center>Progress: 85% (Creating docking results document...)</center></b>")
        try:
            doc = self.create_docking_document(affinity_vals, results_dir, self.mode, self.method)
        except Exception:
            self.finishedSignal.emit([f"❌ Error creating docking document"])
            return
        
        self.toastSignal.emit([2000, "Docking results document created."])

        self.progressBarSignal.emit(99)
        self.progressLabelSignal.emit("<b><center>Progress: 99% (Creating plots...)</center></b>")
        plots_dir = results_dir + "/Plots"
        if not os.path.exists(str(plots_dir)):
            os.mkdir(str(plots_dir))

        all_dicts = [entry for sublist in affinity_vals for entry in sublist]

        receptor_names = [d['Target'] for d in all_dicts]
        binding_energies = [d['Binding Free Energy (kcal/mol)'] for d in all_dicts]
        receptor_energies = list(zip(receptor_names, binding_energies))
        be_plot_path = os.path.join(plots_dir, "Binding_Energy_Distribution.png")
        self.create_binding_energy_plot(receptor_energies, be_plot_path, dpi=self.dpi)
        relative_path = os.path.relpath(be_plot_path)
        self.toastSignal.emit([2000, f"Binding energy distribution plot created at {relative_path}"])
        self.toastSignal.emit([3000, "Docking processing completed successfully!"])
        self.progressLabelSignal.emit("<b><center>Progress: 100% (Completed)</center></b>")
        
        self.finishedSignal.emit(["DONE", self.mode])