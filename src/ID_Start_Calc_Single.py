import os
import re
import csv
import sys
import math
import shutil
import tempfile
import ID_utils
import pyKVFinder
import numpy as np
from glob import glob
import seaborn as sns
from rdkit import Chem
from docx import Document
from rdkit.Chem import QED
from pdbfixer import PDBFixer
from datetime import datetime
from openmm.app import PDBFile
from operator import itemgetter
import matplotlib.pyplot as plt
from docx.shared import Pt, RGBColor
from cryptography.fernet import Fernet
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5.QtCore import QThread, pyqtSignal, QProcess, QEventLoop

LIGAND_TAGS = [' UNL ', ' UNK ', ' LIG ', ' <0> ', ' DRG ', ' INH ', ' NAG ', ' SO4 ', ' ATP ', ' ADP ', ' AMP ', ' HEM ', ' FMN ', ' FAD ', ' NAD ', ' GDP ', ' GTP ', ' SAM ']

STD_RESIDUES = [' ALA ', ' ARG ', ' ASN ', ' ASP ', ' CYS ', ' GLN ', ' GLU ', ' GLY ', ' HIS ', ' ILE ', ' LEU ', ' LYS ', ' MET ', ' PHE ', ' PRO ', ' SER ', ' THR ', ' TRP ', ' TYR ', ' VAL ']

STD_RESIDUES2 = ['ALA','ARG','ASN','ASP','CYS','GLN','GLU','GLY','HIS','ILE','LEU','LYS','MET','PHE','PRO','SER','THR','TRP','TYR','VAL']

key = b'XaQ5nlmxWQyPf9Q9bMjki5M7yEiXgOmTDRpvtTRWbik='

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    if hasattr(sys, '_MEIPASS'):
        # When bundled by PyInstaller, _MEIPASS contains the temporary extraction path.
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

class singleCalculateS1(QThread):
    requestResidueDialog = pyqtSignal()
    residuesReady = pyqtSignal(str)
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(list)

    def setWD(self, wd):
        self.wd = wd

    def checkVals(self, verifylist):
        self.verifylist = verifylist

    def dockMethod(self, text):
        self.dockingMethod = text

    def findProtein(self):
        targets = {}
        for potentialTargetProtein in glob('*'):
            base = os.path.splitext(potentialTargetProtein)[0]
            ext = os.path.splitext(potentialTargetProtein)[1].lower()
            proteinSize = os.path.getsize(potentialTargetProtein) / 1024
            if ext in ['.pdb', '.gpf', '.pdbqt']:
                if base not in targets:
                    targets[base] = {'pdb': None, 'gpf': None, 'pdbqt': None}
                if ext == '.pdb':
                    if proteinSize > 20:
                        with open(potentialTargetProtein, "r") as f:
                            fileContent = f.read()
                            if any(lig in fileContent for lig in LIGAND_TAGS) and any(res in fileContent for res in STD_RESIDUES):
                                self.toastSignal.emit([2000, "Found Co-crystalized ligand in the protein, Continuing..."])
                                targets[base]['pdb'] = potentialTargetProtein
                            elif not any(lig in fileContent for lig in LIGAND_TAGS) and any(res in fileContent for res in STD_RESIDUES):
                                targets[base]['pdb'] = potentialTargetProtein
                elif ext == '.gpf':
                    targets[base]['gpf'] = potentialTargetProtein
                elif ext == '.pdbqt':
                    if proteinSize > 20:
                        with open(potentialTargetProtein, "r") as f:
                            fileContent = f.read()
                            if any(lig in fileContent for lig in LIGAND_TAGS) and any(res in fileContent for res in STD_RESIDUES):
                                self.toastSignal.emit([2000, "Found Co-crystalized ligand in the protein, Continuing..."])
                                targets[base]['pdbqt'] = potentialTargetProtein
                            elif not any(lig in fileContent for lig in LIGAND_TAGS) and any(res in fileContent for res in STD_RESIDUES):
                                targets[base]['pdbqt'] = potentialTargetProtein

        # Filter: must have either pdb or pdbqt
        filtered = {k: v for k, v in targets.items() if v['pdb'] or v['pdbqt']}
        
        # Choose the one with the most populated fields
        def count_populated(entry):
            return sum(1 for v in entry.values() if v is not None)
        
        if filtered:
            best_key = max(filtered, key=lambda k: count_populated(filtered[k]))
            return {best_key: filtered[best_key]}
        else:
            return {None}
    
    def run(self):
        os.chdir(self.wd)
        has_results = os.path.isdir(os.path.join(self.wd, "Results"))
        has_log = os.path.exists(os.path.join(self.wd, "current.idlog"))

        if (has_results and not has_log):
            
            self.finishedSignal.emit([
                "❌ Results folder detected. Please remove it and try again."
            ])
            return
            
        elif has_log and not has_results:
            self.finishedSignal.emit([
                "❌ Error: Inconsistent state detected. Please check your files and try again."
            ])
            return
        
        elif has_log and has_results:
            
            self.toastSignal.emit([2000, "⏳ Resuming the docking process..."])
            
            self.finishedSignal.emit([
                "✅"
            ])
            return

        targets = self.findProtein()
        proteinBaseName, completeProtein = next(iter(targets.items()))

        if completeProtein['gpf'] and not completeProtein['pdb']:
            self.finishedSignal.emit([ f"❌ The target protein '{proteinBaseName}' was recognized as a GPF file but is missing a corresponding PDB file. Please verify your files and try again."])

        if all(completeProtein[attr] is not None for attr in ['pdb', 'gpf', 'pdbqt']):
            if self.dockingMethod == 'GPU':               
                self.finishedSignal.emit(['✅'])
            else:
                self.finishedSignal.emit([f"❌ The target protein '{proteinBaseName}' has corresponding PDB, GPF and PDBQT file but the docking mode is set to using CPU. Please verify your selection and try again."])
        elif (completeProtein['pdb'] and completeProtein['gpf'] and not completeProtein['pdbqt']) and self.dockingMethod == 'GPU':
            self.finishedSignal.emit(['✅'])
        elif (not completeProtein['pdb'] and not completeProtein['gpf'] and completeProtein['pdbqt']) and self.dockingMethod == 'CPU':
            if len(glob("conf.txt")) != 0:
                self.toastSignal.emit([2000, "Calculating bounding box dimensions\nusing user prepared configuration file..."])
                self.finishedSignal.emit(["Site", completeProtein, "CONF"])

            if any(x == 1 for x in self.verifylist):
                if self.verifylist[0] == 1:
                    self.requestResidueDialog.emit()
                    # 2) block *this* thread in its own QEventLoop until we get the text
                    loop = QEventLoop()
                    def _on_text(t):
                        self._residue_text = t
                        loop.quit()
                    # connect the one‐time slot
                    self.residuesReady.connect(_on_text)
                    loop.exec_()  

                    self.toastSignal.emit([2000, "Calculating bounding box dimensions\nusing user selected residues..."])
                    self.finishedSignal.emit(["Site", completeProtein, self._residue_text])

                else:
                    self.finishedSignal.emit([f"❌ The target protein '{proteinBaseName}' was recognized as a PDBQT file but is missing a corresponding PDB file. Please verify your files and try again."])
            
            else:
                self.toastSignal.emit([2000, "Using default blind box dimensions..."])
                self.finishedSignal.emit(["V1", completeProtein, "BLIND"])

        elif (completeProtein['pdb'] and not completeProtein['pdbqt'] and not completeProtein['gpf']) or (completeProtein['pdb'] and completeProtein['pdbqt'] and not completeProtein['gpf']):

            if any(x == 1 for x in self.verifylist):
                if self.verifylist[0] == 1:
                    if len(glob("conf.txt")) != 0:
                        self.toastSignal.emit([2000, "Calculating bounding box dimensions\nusing user prepared configuration file..."])
                        self.finishedSignal.emit(["Site", completeProtein, "CONF"])
                    else:
                        self.requestResidueDialog.emit()
                        # 2) block *this* thread in its own QEventLoop until we get the text
                        loop = QEventLoop()
                        def _on_text(t):
                            self._residue_text = t
                            loop.quit()
                        # connect the one‐time slot
                        self.residuesReady.connect(_on_text)
                        loop.exec_()  

                    self.toastSignal.emit([2000, "Calculating bounding box dimensions\nusing user selected residues..."])
                    self.finishedSignal.emit(["Site", completeProtein, self._residue_text])
                else:
                    self.toastSignal.emit([2000, "Calculating bounding box dimensions using IDv2's binding site predictor..."])
                    self.finishedSignal.emit(["Predict", completeProtein, "PREDICT"])
            
            else:
                self.toastSignal.emit([2000, "Using default blind box dimensions..."])
                self.finishedSignal.emit(["V1", completeProtein, "BLIND"])
            
        else:
            self.finishedSignal.emit(["❌ Please verify your files and try again."])

class singleCalculateS2(QThread):
    finishedSignal = pyqtSignal(list)

    def completeFile(self, file):
        self.mode = file[0]
        self.protein = file[1]
        self.tag = file[2]

    def IDv1Calc(self, proteinLocation):
        # Your existing file parsing to compute center and size.
        xval, yval, zval = [], [], []
        x, y, z = [], [], []
        with open(proteinLocation, 'r') as f:
            for line in f:
                if line.startswith('ATOM'):
                    xval.append(line[30:39])
                    yval.append(line[38:46])
                    zval.append(line[46:54])
        for i in range(len(xval)):
            x.append(float(xval[i]))
            y.append(float(yval[i]))
            z.append(float(zval[i]))

        cx = float(int((sum(x)/len(x))*1000)) / 1000.0
        lx = float(int(max(x) - min(x)) + 10)
        cy = float(int((sum(y)/len(y))*1000)) / 1000.0
        ly = float(int(max(y) - min(y)) + 10)
        cz = float(int((sum(z)/len(z))*1000)) / 1000.0
        lz = float(int(max(z) - min(z)) + 10)

        return [cx, cy, cz], [lx, ly, lz]
    
    def extract_coordinates(self, filename):
        try:
            coords = {}
            with open(filename, 'r') as file:
                for line in file:
                    # Only consider lines with an equals sign
                    if '=' in line:
                        key, value = line.split('=', 1)
                        coords[key.strip()] = value.strip()

            # Retrieve center and size values from the dictionary
            cx = coords.get('center_x')
            cy = coords.get('center_y')
            cz = coords.get('center_z')
            sx = coords.get('size_x')
            sy = coords.get('size_y')
            sz = coords.get('size_z')
            return [cx, cy, cz, sx, sy, sz]
        except:
            return ["❌ No suitable configuration file was found. Please review your files and try again."]

    def run(self):
        if self.mode == 'Site':
            if self.tag == "CONF":
                conf = glob("conf.txt")[0]
                config = self.extract_coordinates(conf)
                if len(config) > 1:
                    center = [config[0], config[1], config[2]]
                    size = [config[3], config[4], config[5]]
                    self.finishedSignal.emit([center, size, self.protein])
                else:
                    self.finishedSignal.emit(config)

            else:
                specificResidues = self.tag.replace(' ', '').split(',')
                resDict = ID_utils.separateSingleResults(specificResidues)
                #Get results as {'A': ['256', '299', '302'], 'B': ['295', '315']} 
                maxRes, minRes = ID_utils.getMaxMinSingle(resDict)
                # Gives (256, 'A'), (310, 'B')
                if self.protein['pdbqt'] and not self.protein['pdb']:
                    first_coords = ID_utils.get_residue_coords(self.protein['pdbqt'], minRes[1], minRes[0])
                    last_coords = ID_utils.get_residue_coords(self.protein['pdbqt'], maxRes[1], maxRes[0])
                else:
                    first_coords = ID_utils.get_residue_coords(self.protein['pdb'], minRes[1], minRes[0])
                    last_coords = ID_utils.get_residue_coords(self.protein['pdb'], maxRes[1], maxRes[0])

                center, size = ID_utils.calculate_bounding_box(first_coords, last_coords)
                self.finishedSignal.emit([center, size, self.protein])

        elif self.mode == 'Predict':
            pdb_path = self.protein.get('pdb')
            try:
                kv_res = pyKVFinder.run_workflow(pdb_path)
            except Exception:
                self.finishedSignal.emit(["❌ No cavities found by IDv2!"])
                return

            vols = kv_res.volume
            cid  = max(vols, key=vols.get)
            unfiltered = kv_res.residues.get(cid, [])
            if not unfiltered:
                self.finishedSignal.emit(["❌ No residues for largest cavity."])
                return
            
            res_list = [entry for entry in unfiltered if entry[2] in STD_RESIDUES2]
            f_res, l_res = res_list[0], res_list[-1]
            c1 = ID_utils.get_residue_coords(self.protein['pdb'], f_res[1], int(f_res[0]))
            c2 = ID_utils.get_residue_coords(self.protein['pdb'], l_res[1], int(l_res[0]))

            center, size = ID_utils.calculate_bounding_box(c1, c2)
            self.finishedSignal.emit([center, size, self.protein])
            return
    
        else:
            # Fallback using default (blind) box dimensions.
            if self.protein['pdbqt'] and not self.protein['pdb']:
                self.proteinLocation = self.protein['pdbqt'] #PDBQT File
                center, size = self.IDv1Calc(self.proteinLocation)
            else:
                self.proteinLocation = self.protein['pdb'] #PDB File
                center, size = self.IDv1Calc(self.proteinLocation)
            
            self.finishedSignal.emit([center, size, self.protein])

class singleCalculateS3(QThread):
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(str)

    def allThings(self, text):        
        self.center = text[0]
        self.size = text[1]
        self.protein = text[2]

    def dockMethod(self, method):
        self.method = method

    def fixPDB(self, checkFix):
        self.checkFix = checkFix

    def fixPdb(self, proteinFile):
        try:
            # Prepare a temporary filename for the fixed version.
            base_name = os.path.splitext(os.path.basename(proteinFile))[0]
            temp_fixed_file = os.path.join(os.path.dirname(proteinFile), f"{base_name}_temp_fixed.pdb")
            
            # Fix the PDB file using PDBFixer.
            fixer = PDBFixer(filename=proteinFile)
            fixer.findMissingResidues()
            fixer.findNonstandardResidues()
            fixer.replaceNonstandardResidues()
            fixer.removeHeterogens(True)
            fixer.findMissingAtoms()
            fixer.addMissingAtoms()
            
            # Write the fixed structure to a temporary file.
            with open(temp_fixed_file, 'w') as f:
                PDBFile.writeFile(fixer.topology, fixer.positions, f, keepIds=True)
            
            # Back up the original file into a backup folder.
            backup_dir = os.path.join(os.path.dirname(proteinFile), "Backup")
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)
            # Here we can rename the original file by appending a suffix.
            backup_file = os.path.join(backup_dir, f"{base_name}_original.pdb")
            shutil.move(proteinFile, backup_file)
            
            # Rename the temporary fixed file to the original filename.
            os.rename(temp_fixed_file, proteinFile)
            
            return "Done"
        
        except Exception:
            return "❌ Unable to fix the PDB, please review your files and try again."
        
    def IDv1Calc(self, proteinLocation):
        # Your existing file parsing to compute center and size.
        xval, yval, zval = [], [], []
        x, y, z = [], [], []
        with open(proteinLocation, 'r') as f:
            for line in f:
                if line.startswith('ATOM'):
                    xval.append(line[30:39])
                    yval.append(line[38:46])
                    zval.append(line[46:54])
        for i in range(len(xval)):
            x.append(float(xval[i]))
            y.append(float(yval[i]))
            z.append(float(zval[i]))

        cx = float(int((sum(x)/len(x))*1000)) / 1000.0
        lx = float(int(max(x) - min(x)) + 10)
        cy = float(int((sum(y)/len(y))*1000)) / 1000.0
        ly = float(int(max(y) - min(y)) + 10)
        cz = float(int((sum(z)/len(z))*1000)) / 1000.0
        lz = float(int(max(z) - min(z)) + 10)

        return [cx, cy, cz], [lx, ly, lz]
        
    def createConfig(self, pdbqtFile, center, size):
        # Write the configuration file in the required format.
        with open("conf.txt", 'w') as f:
            f.write(f"receptor = {pdbqtFile}\n\n")
            f.write(f"center_x = {center[0]:.2f}\n")
            f.write(f"center_y = {center[1]:.2f}\n")
            f.write(f"center_z = {center[2]:.2f}\n\n")
            f.write(f"size_x = {int(size[0])}\n")
            f.write(f"size_y = {int(size[1])}\n")
            f.write(f"size_z = {int(size[2])}\n")

    def convertPDB(self, name, proteinFile):
        if self.method == "GPU":
            gpfFile = name + ".gpf"
            pdbqtFile = name + ".pdbqt"
            ar = [
                "--read_pdb", os.path.abspath(proteinFile),
                "-g", gpfFile, "--write_pdbqt", pdbqtFile,
                "--box_center", str(self.center[0]), str(self.center[1]), str(self.center[2]),
                "--box_size", str(int(self.size[0])), str(int(self.size[1])), str(int(self.size[2]))
            ]
        else:
            pdbqtFile = name + ".pdbqt"
            ar = [
                "--read_pdb", os.path.abspath(proteinFile),
                "-p", pdbqtFile,
                "--box_center", str(self.center[0]), str(self.center[1]), str(self.center[2]),
                "--box_size", str(int(self.size[0])), str(int(self.size[1])), str(int(self.size[2]))
            ]

        # Validate executable path
        primary_exe = resource_path("mk_prepare_receptor.exe")
        if not os.path.exists(primary_exe):
            return f"❌ Executable not found: {primary_exe}"

        # Run mk_prepare_receptor
        self.process.start(primary_exe, ar)
        self.process.waitForFinished(-1)
        output = self.process.readAllStandardOutput().data().decode()

        pdbqtName = name + ".pdbqt"
        if "error" in output.lower() or "failed" in output.lower():
            if self.method == "GPU":
                return "❌ Could not prepare the receptor molecule for the selected docking method"
            # Fallback to OpenBabel
            inputform = '-ipdb'
            alt_exe = "obabel"
            ar = [inputform, proteinFile, "-opdbqt", "-O", pdbqtName, "-xr", "-xc", "-xn", "--partialcharge", "gasteiger"]
            
            self.process.start(alt_exe, ar)
            self.process.waitForFinished(-1)
            output2 = self.process.readAllStandardOutput().data().decode()
            
            if "error" in output2.lower():
                return f"❌ Failed with OpenBabel: {output2}"
            else:
                self.createConfig(pdbqtName, self.center, self.size)
                return "Done"
        else:
            if self.method != "GPU":
                self.createConfig(pdbqtName, self.center, self.size)
            return "Done"
        
    def run(self):
        if self.protein['pdbqt'] and self.protein['pdb'] == None:
            self.createConfig(self.protein['pdbqt'], self.center, self.size)
            self.finishedSignal.emit("✅")
        elif self.protein['pdb']:
            with open(self.protein['pdb'], 'r') as f:
                lines = f.readlines()

            fixedPDB = []
            for line in lines:
                if not line.startswith(('ATOM  ', 'HETATM')):
                    continue
                resname = line[17:20].strip()
                if resname == 'HOH':
                    continue
                if resname not in STD_RESIDUES2:
                    continue
                
                fixedPDB.append(line)

            # Write back out
            with open(self.protein['pdb'], 'w') as f:
                f.writelines(fixedPDB)

            def find_altlocs(pdb_path):
                """
                Scan a PDB file and report any altLoc codes found.
                """
                altlocs = set()
                with open(pdb_path) as f:
                    for line in f:
                        if line.startswith(("ATOM  ", "HETATM")):
                            code = line[16]              # column 17 in 0-based string indexing
                            if code not in (" ", ""):
                                altlocs.add(code)
                return sorted(altlocs)


            def strip_altlocs(pdb_path, output_path, keep="A"):
                """
                Write a new PDB with only the desired altLoc record (default “A”),
                or all atoms that have no altLoc.
                """
                with open(pdb_path) as src, open(output_path, "w") as dst:
                    for line in src:
                        if line.startswith(("ATOM  ", "HETATM")):
                            code = line[16]
                            # keep if either no altLoc or exactly the one we want
                            if code in (" ", "") or code == keep:
                                # blank out column 17 so downstream tools see no alternate locs
                                dst.write(line[:16] + " " + line[17:])
                        else:
                            dst.write(line)

            pdb_path = self.protein['pdb']
            altlocs = find_altlocs(pdb_path)
            if altlocs:
                dirn, base = os.path.split(pdb_path)
                fd, tmp_path = tempfile.mkstemp(prefix=base, dir=dirn, text=True)
                os.close(fd)
                strip_altlocs(pdb_path, tmp_path, keep=altlocs[0])
                os.replace(tmp_path, pdb_path)

            self.process = QProcess()
            self.process.setProcessChannelMode(QProcess.MergedChannels)
            primary_exe = resource_path("mk_prepare_receptor.exe")
            minimal_args = ["--read_pdb", os.path.abspath(self.protein['pdb'])]
            self.process.start(primary_exe, minimal_args)
            self.process.waitForFinished(-1)
            output = self.process.readAllStandardOutput().data().decode()

            if (("error" in output.lower() or "failed" in output.lower()) and (self.method == "GPU" or self.checkFix == 1)):
                self.toastSignal.emit([2000, "Fixing and conversion in progress..."])
                testTxt = self.fixPdb(self.protein['pdb'])
                
            else:
                self.toastSignal.emit([2000, "Conversion in progress..."])
                testTxt = "DoneWithoutFix"
            
            if testTxt == "Done":
                name = self.protein['pdb'].split('.')[0]
                proteinFixed = self.protein['pdb']
                testConv = self.convertPDB(name, proteinFixed)              
                self.finishedSignal.emit(testConv)

            elif testTxt == "DoneWithoutFix":
                name = self.protein['pdb'].split('.')[0]
                testConv = self.convertPDB(name, self.protein['pdb'])               
                self.finishedSignal.emit(testConv)

            else:               
                self.finishedSignal.emit(testTxt)

        else:
            self.finishedSignal.emit("❌ No target file could be found, please check your working directory and try again.")

class singleCalculateS4(QThread):
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(str)
        
    def setWD(self, wd):
        self.wd = wd

    def extractLibs(self, libraries):
        cwd = os.getcwd()
        libmulti = " "
        libsingle = " "
        ligCount = 0
        if len(libraries) > 1:
            libmulti = "Ligand_libraries_found"
            for library in libraries:
                if library.endswith(".mol2"):
                    with open(library, mode='r', encoding='utf-8') as f:
                        text = f.read()
                    d = "@<TRIPOS>MOLECULE"
                    splitted = [d+e for e in text.split(d) if e]
                    for i in splitted:
                        ligCount += 1
                        test = i.splitlines()
                        name = test[1]
                        outputfile =  name+".mol2"
                        with open(outputfile,'wb') as u:
                            file = "\n".join(test).decode('utf-8')
                            u.write(file)

                elif library.endswith(".sdf"):
                    with open(library, mode='r', encoding='utf-8')  as f:
                        text = f.read()
                    arr = text.split("$$$$")[:-1]
                    for file in arr:
                        ligCount += 1
                        if file == arr[0]:
                            name = file.splitlines()[0]
                        elif file == arr[len(arr)-1]:
                            name = file.splitlines()[1]
                        else:
                            name = file.splitlines()[1]
                        with open(os.path.join(cwd, name+".sdf"),'wb') as u:
                            file = "".join(file[1:len(file)-1])
                            u.write(file.encode('utf-8'))

        elif len(libraries) == 1:
            libsingle = "Ligand_library_found"
            libraries = libraries[0]
            
            if libraries.endswith(".mol2"):
                with open(libraries, mode='r', encoding='utf-8') as f:
                    text = f.read()
                d = "@<TRIPOS>MOLECULE"
                splitted = [d+e for e in text.split(d) if e]
                for i in splitted:
                    ligCount += 1
                    test = i.splitlines()
                    name = test[1]
                    outputfile =  name+".mol2"
                    with open(outputfile,'wb') as u:
                        file = "\n".join(test).decode('utf-8')
                        u.write(file)

            elif libraries.endswith(".sdf"):
                with open(libraries, mode='r', encoding='utf-8') as f:
                    text = f.read()
                arr = text.split("$$$$")[:-1]
                for file in arr:
                    ligCount += 1
                    if file == arr[0]:
                        name = file.splitlines()[0]
                    elif file == arr[len(arr)-1]:
                        name = file.splitlines()[1]
                    else:
                        name = file.splitlines()[1]
                    with open(os.path.join(cwd, name+".sdf"),'wb') as u:
                        file = "".join(file[1:len(file)-1])
                        u.write(file.encode('utf-8'))

        if libmulti != " ":
            lib = libmulti
            if not os.path.exists(libmulti):
                os.makedirs(libmulti)
                for library in libraries:
                    shutil.move(os.path.join(cwd, library), libmulti)
        else:
            pass
        
        if libsingle != " ":
            lib = libsingle
            if not os.path.exists(libsingle):
                os.makedirs(libsingle)
                shutil.move(os.path.join(cwd, libraries), libsingle)
        else:
            pass
        
        self.toastSignal.emit([2000, f"{ligCount} ligands extracted."])
        return lib
    
    def findLigs(self):       
        ligandFiles = list()
        libraryFiles = list()
        potentialLigands = glob('*.pdb')+ glob('*.mol2') + glob('*.sdf') + glob('*.mol') + glob('*.pdbqt')
        for ligFile in potentialLigands:
            with open(ligFile, "r") as f:
                fileContent = f.read()
            if any(res in fileContent for res in STD_RESIDUES) or "BOX" in fileContent:
                pass
            else:
                ext = os.path.splitext(ligFile)[1].lower()
                if ext == '.sdf':
                    # SDF multi‐molecule delimiter is $$$$
                    if fileContent.count('$$$$') > 1:
                        libraryFiles.append(ligFile)
                    else:
                        ligandFiles.append(ligFile)

                elif ext == '.mol2':
                    # MOL2 multi‐molecule marker is @<TRIPOS>MOLECULE
                    if fileContent.count('@<TRIPOS>MOLECULE') > 1:
                        libraryFiles.append(ligFile)
                    else:
                        ligandFiles.append(ligFile)

                else:
                    ligandFiles.append(ligFile)

        return ligandFiles, libraryFiles
    
    def convertLig(self, ligFile):
        # Ensure the ligand file is an absolute path.
        ligFile = os.path.abspath(ligFile)
        
        if ligFile.endswith(".pdbqt"):
            pass
        else:       
            # Build an absolute path to the executable.
            cd = "obabel"
            format = ligFile.split('.')[1]
            inputform = "-i"+format
            tem = re.sub('\..*', '.pdbqt', ligFile)
            lig = tem
            ar = [inputform, ligFile, "-opdbqt", "-O", lig, "--partialcharge gasteiger", "-h"]
            
            self.process = QProcess()
            self.process.execute(cd, ar)
            self.process.waitForFinished(-1)

    def run(self):
        os.chdir(self.wd)
        ligandFiles, libraryFiles = self.findLigs()
        if libraryFiles:
            self.extractLibs(libraryFiles)
        ligandFiles, libraryFiles = self.findLigs()
        if ligandFiles:
            for ligandFile in ligandFiles:
                self.convertLig(ligandFile)
                
            if len(ligandFiles) > 1:  
                self.toastSignal.emit([2000, f"{len(ligandFiles)} candidates ready for docking analysis!"])             
                self.finishedSignal.emit("✅")
            else:               
                self.toastSignal.emit([2000, f"Ligand {ligandFiles[0]} is ready for docking analysis."])
                self.finishedSignal.emit("✅")
        else:           
            self.finishedSignal.emit("❌ No ligands found! Please check your files and try again.")

class startSingleTargetDocking(QThread):
    toastSignal = pyqtSignal(list)
    finishedSignal = pyqtSignal(str)
    progressBarSignal = pyqtSignal(int)
    progressLabelSignal = pyqtSignal(str)

    def setWD(self, wd):
        self.wd = wd

    def checkVals(self, verifylist):
        self.resume = True if verifylist[0] else False
        self.logging = True if verifylist[1] else False

    def dockMethod(self, text):
        self.dockingMethod = text

    def setNConformers(self, text):
        self.conformers = text

    def resumeDocking(self, dir):
        log_path = os.path.join(dir, "current.idlog")
        
        if not os.path.exists(log_path):
            return []
        
        try:
            with open(log_path, "rb") as file:
                encrypted = file.read()
            decrypted = Fernet(key).decrypt(encrypted).decode("utf-8")
            log_entries = decrypted.strip().split('\n')
            
            ligands = set()
            for entry in log_entries:
                try:
                    ligands.add(entry)
                except ValueError:
                    continue
            
            completed = []
            for ligand in ligands:  # Split from "target_path::ligand_path"
                completed.append(ligand)
            return completed
        
        except Exception:
            self.finishedSignal.emit(f"❌ Error parsing log.")
            return []

    def process_docked_file(self, input_file, output_file):
        """
        Process an AutoDock DLG file to create a cleaned PDBQT file with modified energy remarks.
        This method removes the 'DOCKED:' prefix and replaces a 'USER' line (containing the energy)
        with a 'REMARK' line.
        """
        energy_identifier = "Estimated Free Energy of Binding"
        with open(input_file, 'r') as infile:
            lines = infile.readlines()
        
        processed_lines = []
        for line in lines:
            if line.startswith('DOCKED:'):
                clean_line = line[7:].lstrip()  # Remove 'DOCKED:' and any leading whitespace
                if clean_line.startswith('USER') and energy_identifier in clean_line:
                    # Replace 'USER' with 'REMARK' and rename the energy identifier
                    modified_line = clean_line.replace('USER', 'REMARK', 1)
                    modified_line = modified_line.replace(energy_identifier, 'Binding energy', 1)
                    processed_lines.append(modified_line)
                elif not clean_line.startswith('USER'):
                    processed_lines.append(clean_line)
        with open(output_file, 'w', newline='') as outfile:
            outfile.writelines(processed_lines)

    def findLigs(self, pdbqtFiles):
        """
        Filters the list of pdbqt files:
         - Skip files containing any standard residue (STD_RESIDUES).
         - Otherwise, consider it as a ligand file.
        """
        ligandFiles = []
        for ligFile in pdbqtFiles:
            with open(ligFile, "r") as f:
                fileContent = f.read()
                if any(res in fileContent for res in STD_RESIDUES):
                    continue
                else:
                    ligandFiles.append(ligFile)

        return ligandFiles

    @staticmethod
    def diff_sorted_lists(total_ligands, completed_ligands):
        """Find the last common ligand between the total and completed lists.
        Returns the sublist of `total_ligands` from the last common element onward."""
        i = j = 0
        last_common_index = -1  # Initialize to -1 to include the first element if needed

        while i < len(total_ligands) and j < len(completed_ligands):
            if total_ligands[i] == completed_ligands[j]:
                last_common_index = i  # Update to current index in total_ligands
                i += 1
                j += 1
            elif total_ligands[i] < completed_ligands[j]:
                i += 1
            else:
                j += 1

        # Return from the last_common_index (inclusive) to the end of total_ligands
        return total_ligands[last_common_index + 1:] if last_common_index != -1 else total_ligands.copy()

    def _save_intermediate_log(self):
        try:
            log_entries = [f"{ligand}" for ligand in self.log]
            plain_text = "\n".join(log_entries).encode("utf-8")
            fernet = Fernet(key)
            encrypted = fernet.encrypt(plain_text)
            
            # Atomic write
            temp_path = self.temp_log_path + ".tmp"
            with open(temp_path, "wb") as temp_file:
                temp_file.write(encrypted)
            os.replace(temp_path, self.temp_log_path)
        except Exception: 
            self.finishedSignal.emit(f"❌ Error saving log.")
            return
        
    def normalize(self, p):
        # normpath will collapse mixed slashes, normcase lower‑cases drive letters on Windows
        return os.path.normcase(os.path.normpath(p))
        
    def run(self):
        self.log = []
        self.temp_log_path = os.path.join(self.wd, "current.idlog")
        ligs = glob(os.path.join(self.wd, "*.pdbqt"))
        ligFiles = self.findLigs(ligs)
        total_ligands = len(ligFiles)
        if total_ligands == 0:
            self.finishedSignal.emit("❌ No ligand files found (.pdbqt)!")
            return
        resultFolder = os.path.join(self.wd, "Results")
        if not os.path.exists(resultFolder):
            os.mkdir(resultFolder)
        sorted_total_lig_files = sorted(ligFiles, key=os.path.getsize)
        
        # Apply resume filtering if needed.
        if self.resume:
            # 1) Read and normalize your completed list
            completed_pdbqt_files = self.resumeDocking(self.wd)
            completed_norm = {self.normalize(p) for p in completed_pdbqt_files}

            # 2) Sanity check: make sure nothing weird sneaked into the log
            total_norm = {self.normalize(p) for p in sorted_total_lig_files}
            if not completed_norm.issubset(total_norm):
                self.finishedSignal.emit(
                    "❌ IDLOG file compounds do not match current compounds."
                )
                return

            # 3) Find the last‐docked index (highest i where ligand ∈ log)
            last_idx = -1
            for i, lig in enumerate(sorted_total_lig_files):
                if self.normalize(lig) in completed_norm:
                    last_idx = i

            # 4) Slice **from** that index, so you re‐dock the last one
            if last_idx >= 0:
                sorted_lig_files = sorted_total_lig_files[last_idx:]
            else:
                # nothing completed yet, run them all
                sorted_lig_files = sorted_total_lig_files
        else:
            sorted_lig_files = sorted_total_lig_files

        lig_count = len(sorted_lig_files)
        self.progressLabelSignal.emit(f"<center><b>Progress: 0% (0/{lig_count}) ligands</b>")

        if self.dockingMethod == "AutoDock-GPU":            
            cueMethod = "GPU"
            gpfFiles = glob(os.path.join(self.wd, "*.gpf"))
            if not gpfFiles:
                self.finishedSignal.emit("❌ No GPF found! Please check your files.")
                return
            try:
                autogrid_src = resource_path("autogrid4.exe")
                dlls = [resource_path("cyggcc_s-seh-1.dll"), resource_path("cyggomp-1.dll"), resource_path("cygstdc++-6.dll"), resource_path("cygwin1.dll")]
                DLLS = ["cyggcc_s-seh-1.dll", "cyggomp-1.dll", "cygstdc++-6.dll", "cygwin1.dll"]
                shutil.copy(autogrid_src, self.wd)
                for dll in dlls:
                    shutil.copy(dll, self.wd)

            except Exception:
                self.finishedSignal.emit(f"❌ Failed to complete task.")
                return
            
            if gpfFiles:
                protein_file = gpfFiles[0]
                protein_base = os.path.splitext(protein_file)[0]
                grid_file = f"{protein_base}.maps.fld"
                sorted_lig_files = [lig for lig in sorted_lig_files if os.path.splitext(lig)[0] != protein_base]
            
            if not os.path.exists(grid_file):
                self.toastSignal.emit([5000, "Initializing grid parameters… this may take a while, please be patient."])
                try:
                    autogrid_exe = os.path.join(self.wd, "autogrid4.exe")
                    cmd = f'"{autogrid_exe}" -p "{protein_file}"'
                    self.process = QProcess()
                    self.process.execute(cmd)
                    self.process.waitForFinished(-1)
                    os.remove(autogrid_exe)
                    for DLL in DLLS:
                        os.remove(os.path.join(self.wd, DLL))

                except Exception:
                    self.finishedSignal.emit(f"❌ Grid setup failed")
                    return
            
            autodock_gpu_exe = resource_path("AutoDock-GPU.exe")
            self.process = QProcess()
            for j, lig in enumerate(sorted_lig_files):                
                rotatable_count = 0
                with open(lig, 'r') as lig_f:
                    for line in lig_f:
                        # each REMARK line for a rotatable bond starts like "REMARK    <num>  A    between atoms:"
                        if line.startswith("REMARK") and "between atoms:" in line:
                            rotatable_count += 1

                if rotatable_count > 32:
                    # notify & skip
                    self.toastSignal.emit([
                        2000,
                        f"Skipping {lig_name}: {rotatable_count} rotatable bonds (>32) not supported on GPU."
                    ])
                    # update progress bar anyway, so user sees you're moving on
                    progress = int((j + 1) / lig_count * 75)
                    self.progressBarSignal.emit(progress)
                    self.progressLabelSignal.emit(
                        f"<center><b>Progress: {progress}% ({j+1}/{lig_count}) ligands</b>"
                    )
                    continue
                # Docking command
                lig_name = os.path.splitext(lig)[0]
                out_base = os.path.join(self.wd, f"{lig_name}_out")
                cmd = f'"{autodock_gpu_exe}" --lfile "{lig}" --ffile "{grid_file}" --nrun {self.conformers} --xmloutput 0 --N "{out_base}"'
                self.process.execute(cmd)
                self.process.waitForFinished(-1)
                
                # Log and save
                dlg_path = f"{out_base}.dlg"
                pdbqt_out = f"{out_base}.pdbqt"
                self.process_docked_file(dlg_path, pdbqt_out)

                self.log.append(os.path.abspath(lig))
                self._save_intermediate_log()
                
                # Calculate progress (based on totalDockingSteps)
                progress = int((j + 1) / lig_count * 75)  # 75% cap
                self.progressBarSignal.emit(progress)
                self.progressLabelSignal.emit(f"<center><b>Progress: {progress}% ({j+1}/{lig_count}) ligands</b>")

        else:  # AutoDock Vina/QuickVina-W
            config_path = os.path.join(self.wd, "conf.txt")
            if not os.path.exists(config_path):
                self.finishedSignal.emit("❌ No configuration file (conf.txt) found!")
                return
            
            # Rename ligands to avoid spaces/dashes
            renamed_ligands = []
            for idx, lig in enumerate(sorted_lig_files):
                baseLig = os.path.basename(lig)
                new_name = baseLig.replace(" ", "_").replace("-", "")
                if new_name != baseLig:
                    os.rename(lig, new_name)
                    renamed_ligands.append(new_name)
                else:
                    renamed_ligands.append(lig)
            sorted_lig_files = renamed_ligands

            if self.dockingMethod == "AutoDock Vina":
                exe_path = resource_path("vina.exe")
                cueMethod = "VINA"
            else:
                exe_path = resource_path("qvinaw.exe")
                cueMethod = "QVW"

            for idx, lig in enumerate(sorted_lig_files):
                lig_name = os.path.splitext(lig)[0]
                out_file = f"{lig_name}_out.pdbqt"
                log_file = f"{lig_name}_log.log"

                cmd = f'"{exe_path}" --config "{config_path}" --ligand "{lig}" --num_modes {self.conformers} --out "{out_file}"'
                if self.logging and self.dockingMethod == "QuickVina-W":
                    cmd += f' --log "{log_file}"'

                self.process = QProcess()
                self.process.execute(cmd)
                self.process.waitForFinished(-1)

                # Log and update progress (75% scaling)
                self.log.append(lig)
                self._save_intermediate_log()
                progress = int((idx + 1) / lig_count * 75)  # 75% cap
                self.progressBarSignal.emit(progress)
                self.progressLabelSignal.emit(f"<center><b>Progress: {progress}% ({idx+1}/{lig_count}) ligands</b>")

        cue = f'DONE-{cueMethod}'

        self.finishedSignal.emit(cue)

        if self.log:
            current_time = datetime.now()
            formatted_time = current_time.strftime("%Y-%m-%d_%H-%M-%S")
            final_log_path = os.path.join(self.wd, f"{formatted_time}.idlog")
            try:
                os.rename(self.temp_log_path, final_log_path)
            except FileNotFoundError:
                # Fallback if temp file is missing (first run)
                with open(final_log_path, "wb") as file:
                    fernet = Fernet(key)
                    file.write(fernet.encrypt('\n'.join([f"{l}" for l in self.log]).encode()))

class continueSingleID(QThread):
    requestInput = pyqtSignal(int)
    progressBarSignal = pyqtSignal(int)
    progressLabelSignal = pyqtSignal(str)
    finishedSignal = pyqtSignal(list)
    toastSignal = pyqtSignal(list)
    inputReceived = False

    def __init__(self):
        super().__init__()
        self.top_hits = None  # Will hold the user's answer
        self.loop = QEventLoop()

    def setInputResult(self, value):
        # Called from the main thread once the input is retrieved.
        self.top_hits = value
        self.inputReceived = True
        # Exit the local event loop so run() can continue.
        self.loop.quit()

    def setWD(self, wd):
        # Set self.wd as the parent folder where the Results folder (and its subfolders) will be created.
        self.wd = wd
        
    def setDPI(self, dpi):
        self.dpi = int(dpi)

    def setBlind(self, blind):
        self.blind = blind

    def setBinding(self, binding):
        self.binding = binding

    def setSiteSpecific(self, siteSpecific):
        self.siteSpecific = siteSpecific

    def dockMethod(self, method):
        # Set the method description along with citation numbers.
        if method == "DONE-VINA":
            self.method = "AutoDock Vina [2]"
        elif method == "DONE-QVW":
            self.method = "QuickVina-W [2] (modified AutoDock Vina [3])"
        elif method == "DONE-GPU":
            self.method = "AutoDock-GPU [2]"

    def setLogging(self, log):
        self.logging = log

    def findTarget(self):
        if self.method == "AutoDock-GPU [2]":
            proteinName = glob("*.maps.fld")
            return proteinName
        else:
            with open("conf.txt", 'r') as f:
                for line in f:
                    if line.startswith("receptor"):
                        return line.split("=")[1].strip()
                    
    def moveTarget(self, receptorName):
        target_dir = self.wd + f"/{receptorName} related files"
        if not os.path.exists(str(target_dir)):
            os.mkdir(str(target_dir))
        for pattern in [f"{receptorName}*", "boron-silicon-atom_par.dat"]:
            for file in glob(pattern):
                if file not in f"{receptorName} related files":
                    shutil.move(file, str(target_dir) + os.sep + os.path.basename(file))        
                    
    def human_format(num):
        suffixes = [
            (1e12, 'Trillion'),
            (1e9, 'Billion'),
            (1e6, 'Million'),
            (1e3, 'Thousand'),
        ]
        for val, suffix in suffixes:
            if num >= val:
                formatted_num = num / val
                # Use .3g to format up to 3 significant figures
                return f"{formatted_num:.3g} {suffix}"
        return str(num)

    def create_docking_document(self, receptor_name, ligands, results_dir, mode, top_hits, dockMethod):
        #receptor_name = full with format
        #results_dir = complete results path
        #mode = vs or docking
        #top_hits = selected top hits
        #dockMethod = "AutoDock Vina [2]", "QuickVina-W [2] (modified AutoDock Vina [3])" OR "AutoDock-GPU [2]"
        #Ligands contains a list of dictionaries with all ligand information.
        # Example.:'Name': LigandName, 'Binding Energy (kcal/mol)': -6.0, 'pKi': 4.401, 'Ligand Efficiency (kcal/mol/non-H atom)': 0.051, 'Torsional Energy': 1.245, 'Molecular Weight': 16.0, 'HBD': 0, 'HBA': 0, 'LogP': 0.6, 'Lipinski Violations': '0', 'Lipinski Compliant': 'True'

        document = Document()
        styles = document.styles

        style_one = styles.add_style('Heading One', WD_STYLE_TYPE.PARAGRAPH)
        style_one.base_style = styles['Heading 1']
        font = style_one.font
        font.name = 'Times New Roman'
        font.size = Pt(16)
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)

        style_three = styles.add_style('Heading Three', WD_STYLE_TYPE.PARAGRAPH)
        style_three.base_style = styles['Heading 2']
        font = style_three.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = False
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)     

        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            methodDescription = "a hybrid scoring function (empirical + knowledge-based) in docking calculations"
            receptor = receptor_name.split('.')[0]
        elif dockMethod == "AutoDock Vina [2]":
            methodDescription = "a scoring function based on the Vina algorithm"
            receptor = receptor_name.split('.')[0]
        elif dockMethod == "AutoDock-GPU [2]":
            receptor = receptor_name[0].split('.')[0]
            methodDescription = "a scoring function based on the AutoDock4 algorithm"

        if self.blind and not self.binding:
            approach = "a blind search space for the ligand"
        if self.blind and self.binding:
            approach = "a predicted search space for the ligand"
        if self.siteSpecific:
            approach = "a site-specific search space for the ligand"

        if mode == 'docking':
            modeToWrite = ''
        else:
            modeToWrite = "-based virtual screening"
        
        if len(ligands) == 1:
            ligand = ligands[0]['Name']
            para = document.add_paragraph(f'Molecular docking of \"{ligand}\" with \"{receptor}\"', style = 'Heading One')
            pformat = para.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = document.add_paragraph(f'Molecular docking of \"{ligand}\" with \"{receptor}\" was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
            p.add_run(f"The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}.")
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif len(ligands) == 2:
            ligand1 = ligands[0]['Name']
            ligand2 = ligands[1]['Name']
            para = document.add_paragraph(f'Molecular docking of \"{ligand1}\" and \"{ligand2}\" with \"{receptor}\"', style = 'Heading One')
            pformat = para.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = document.add_paragraph(f'Molecular docking of \"{ligand1}\" and \"{ligand2}\" with \"{receptor}\" was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
            p.add_run(f'The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}s.')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif len(ligands) > 2:
            if mode == 'docking':
                ligLen = len(ligands)
                para = document.add_paragraph(f'Molecular docking of {len(ligands)} compounds with \"{receptor}\"', style = 'Heading One')
                pformat = para.paragraph_format
                pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = document.add_paragraph(f'Molecular docking of {len(ligands)} compounds with \"{receptor}\" was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
                p.add_run(f'The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}s.')
                pformat = p.paragraph_format
                pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            else:
                ligLen = len(ligands)
                if ligLen > 10000:
                    ligLen = self.human_format(len(ligands))
                para = document.add_paragraph(f'Virtual screening of {ligLen} compounds with \"{receptor}\"', style = 'Heading One')
                pformat = para.paragraph_format
                pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = document.add_paragraph(f'Virtual screening of {ligLen} compounds with \"{receptor}\" was performed to predict their binding affinity and detailed interactions. ', style='Heading Three')
                p.add_run(f'The docking was performed using InstaDock-v2, a molecular docking tool that automizes and enhances the entire process of molecular docking{modeToWrite} [1]. The binding affinities between the ligand and protein were calculated using the {dockMethod} program which uses {methodDescription} and {approach}s.')
                pformat = p.paragraph_format
                pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        style_two = styles.add_style('Heading Two', WD_STYLE_TYPE.PARAGRAPH)
        style_two.base_style = styles['Heading 2']
        font = style_two.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)
        paragraph = document.add_paragraph('Materials and Methods', style = 'Heading Two')
        pformat = paragraph.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        p = document.add_paragraph('The ', style='Heading Three')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.add_run('p')
        p.add_run('Ki').italic = True
        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            p.add_run(', the negative decimal logarithm of inhibition constant [4] was calculated from the ∆')
        else:
            p.add_run(', the negative decimal logarithm of inhibition constant [3] was calculated from the ∆')
        p.add_run('G').italic = True
        p.add_run(' parameter while using the following formula:')
        style_six = styles.add_style('Heading Six', WD_STYLE_TYPE.PARAGRAPH)
        style_six.base_style = styles['Heading 6']
        font = style_six.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        font.bold = False
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)
        d = document.add_paragraph('', style='Heading Three')
        dformat = d.paragraph_format
        dformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        d.add_run('∆')
        d.add_run('G').italic = True
        d.add_run(' = RT(Ln ')
        d.add_run('Ki').italic = True
        test = d.add_run('pred')
        test.font.subscript = True
        d.add_run(')')
        run = d.add_run()
        run.add_break()
        d.add_run('Ki').italic = True
        test = d.add_run('pred')
        test.font.subscript = True
        d.add_run(' = e')
        test = d.add_run('(∆')
        test.font.superscript = True
        test = d.add_run('G')
        test.font.superscript = True
        test.italics = True
        test = d.add_run('/RT)')
        test.font.superscript = True
        run = d.add_run()
        run.add_break()
        d.add_run('p')
        d.add_run('Ki').italic = True
        d.add_run(' = -log(')
        d.add_run('Ki').italic = True
        test = d.add_run('pred')
        test.font.subscript = True
        d.add_run(')')
        p = document.add_paragraph('where ∆', style='Heading Three')
        p.add_run('G').italic = True
        p.add_run(' is the binding affinity (kcal mol')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run('), R (gas constant) is 1.98 cal*(mol*K)')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run(', T (room temperature) is 298.15 Kelvin, and ')
        p.add_run('Ki').italic = True
        test = p.add_run('pred')
        test.font.subscript = True
        p.add_run(' is the predicted inhibitory constant.')

        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            p = document.add_paragraph("Ligand efficiency (LE) is a commonly applied parameter for selecting favorable ligands by comparing the values of average binding energy per atom [5]. The following formula was applied to calculate LE:", style='Heading Three')
        else:
            p = document.add_paragraph("Ligand efficiency (LE) is a commonly applied parameter for selecting favorable ligands by comparing the values of average binding energy per atom [4]. The following formula was applied to calculate LE:", style='Heading Three')
        p.add_run("")
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        d = document.add_paragraph('', style='Heading Three')
        dformat = d.paragraph_format
        dformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        d.add_run('LE = -∆')
        d.add_run('G').italic = True
        d.add_run('/N')
        p = document.add_paragraph('where LE is the ligand efficiency (kcal mol', style='Heading Three')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run(' non-H atom')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run('), ∆')
        p.add_run('G').italic = True
        p.add_run(' is binding affinity (kcal mol')
        test = p.add_run('-1')
        test.font.superscript = True
        p.add_run(') and N is the number of non-hydrogen atoms in the ligand.')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            p = document.add_paragraph(
                "Quantitative Estimate of Drug-likeness (QED) assesses how 'drug-like' a molecule is "
                "by combining multiple physicochemical properties [6]. The formula being:",
                style='Heading Three'
            )
        else:
            p = document.add_paragraph(
                "Quantitative Estimate of Drug-likeness (QED) assesses how 'drug-like' a molecule is "
                "by combining multiple physicochemical properties [5]. The formula being:",
                style='Heading Three'
            )

        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # render the QED equation
        qed_eq = document.add_paragraph('', style='Heading Three')
        qed_eq_format = qed_eq.paragraph_format
        qed_eq_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qed_eq_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Construct QED = (d₁ʷ¹ × d₂ʷ² × ... × d₈ʷ⁸)^(1 / ∑wᵢ)
        qed_eq.add_run('QED = (')
        qed_eq.add_run('d')
        qed_eq.add_run('1').font.subscript = True
        qed_eq.add_run('w1').font.superscript = True
        qed_eq.add_run(' x ')
        qed_eq.add_run('d')
        qed_eq.add_run('2').font.subscript = True
        qed_eq.add_run('w2').font.superscript = True
        qed_eq.add_run(' x ')
        qed_eq.add_run('...')
        qed_eq.add_run(' x ')
        qed_eq.add_run('d')
        qed_eq.add_run('8').font.subscript = True
        qed_eq.add_run('w8').font.superscript = True
        qed_eq.add_run(')')
        qed_eq.add_run('1/∑').font.superscript = True
        p = qed_eq.add_run('wi')
        p.font.italic = True
        p.font.superscript = True

        # a brief “where …” description
        p = document.add_paragraph(
            "where each dᵢ is the desirability of property i, and wᵢ is its corresponding weight.",
            style='Heading Three'
        )
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        p = document.add_paragraph('Results', style = 'Heading Two')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if len(ligands) == 1:
            ligand = ligands[0]['Name']
            bEnergy = ligands[0]['Binding Free Energy (kcal/mol)']
            pKi = ligands[0]['pKi']
            lEfficiency = ligands[0]['Ligand Efficiency (kcal/mol/non-H atom)']
            p = document.add_paragraph(f'Compound \"{ligand}\" was subjected to docking analysis, it presented a binding affinity value of {bEnergy} kcal mol', style='Heading Three')
            test = p.add_run('-1')
            test.font.superscript = True
            p.add_run(' and a p')
            p.add_run('Ki').italic = True
            p.add_run(f' value of {pKi} towards the receptor \"{receptor}\". It also possesses a ligand efficiency of {lEfficiency} kcal mol ')
            test = p.add_run('-1')
            test.font.superscript = True
            p.add_run(' non-H atom')
            test = p.add_run('-1')
            test.font.superscript = True
            p.add_run(f'. Please refer to the CSV file given in the central directory for other docking parameters of the compound used in this study. Interaction analysis of the all possible docked conformers of \"{ligand}\" was carried out to investigate their binding pattern and possible interactions towards the \"{receptor}\" binding pocket. ')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif len(ligands) == 2:
            ligand1 = ligands[0]['Name']
            ligand2 = ligands[1]['Name']
            bEnergy1 = ligands[0]['Binding Free Energy (kcal/mol)']
            bEnergy2 = ligands[1]['Binding Free Energy (kcal/mol)']
            p = document.add_paragraph(f'Compound \"{ligand1}\" and \"{ligand2}\" were subjected to docking analysis, they presented a binding affinity value of {bEnergy1} kcal/mol and {bEnergy2} kcal/mol, respectively towards the receptor \"{receptor}\". Please refer to the CSV file given in the central directory for other docking parameters of each compound used in this study. Interaction analysis of the all possible docked conformers of both the compounds was carried out to investigate their binding pattern and possible interactions towards the \"{receptor}\" binding pocket.', style='Heading Three')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        else:
            sorted_energies = sorted([d['Binding Free Energy (kcal/mol)'] for d in ligands])
            lowest = sorted_energies[0]
            highest = sorted_energies[-1]
            best_ligand = min(ligands, key=lambda x: x['Binding Free Energy (kcal/mol)'])
            bestName = best_ligand['Name']
            bestBind = best_ligand['Binding Free Energy (kcal/mol)']
            p = document.add_paragraph(f'All {ligLen} compounds were subjected to docking analysis and presented a binding affinity within the range of {lowest} kcal/mol to {highest} kcal/mol towards \"{receptor}\". Please refer to the CSV file given in the central directory for the binding affinities and other docking parameters of each compound used in this study, where the best binding affinity was observed in the case of \"{bestName}\" as {bestBind} kcal/mol. Interaction analysis of the all possible docked conformers of all compounds was carried out to investigate their binding pattern and possible interactions towards the \"{receptor}\" binding pocket.', style='Heading Three')
            pformat = p.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        p = document.add_paragraph('References', style = 'Heading Two')
        pformat = p.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        style_four = styles.add_style('Heading Four', WD_STYLE_TYPE.PARAGRAPH)
        style_four.base_style = styles['List Number']
        font = style_four.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = False
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)

        c = document.add_paragraph('Mathur Y, Hassan MI (2025). InstaDock-v2: TBD. ', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run(' TBD').italic = True #Journal
        c.add_run(', ')
        c.add_run('TBD').bold = True #Volume
        c.add_run('TBD') #Pages and DOI
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if dockMethod == "QuickVina-W [2] (modified AutoDock Vina [3])":
            c = document.add_paragraph('Hassan NM, Alhossary AA, Mu Y, Kwoh CK (2017). Protein-ligand blind docking using QuickVina-W with inter-process spatio-temporal integration,' , style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Scientific Reports').italic = True
            c.add_run(', ')
            c.add_run('7(1)').bold = True
            c.add_run(': 1-13.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            c = document.add_paragraph('Jerome Eberhardt, Diogo Santos-Martins, Andreas F. Tillack, and Stefano Forli (2021) AutoDock Vina 1.2.0: New Docking Methods, Expanded Force Field, and Python Bindings.', style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Journal of Chemical Information and Modeling').italic = True
            c.add_run(', ')
            c.add_run('61(8)').bold = True
            c.add_run(': 3891-3898.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif dockMethod == "AutoDock Vina [2]":
            c = document.add_paragraph('Jerome Eberhardt, Diogo Santos-Martins, Andreas F. Tillack, and Stefano Forli (2021) AutoDock Vina 1.2.0: New Docking Methods, Expanded Force Field, and Python Bindings.', style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Journal of Chemical Information and Modeling').italic = True
            c.add_run(', ')
            c.add_run('61(8)').bold = True
            c.add_run(': 3891-3898.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        elif dockMethod == "AutoDock-GPU [2]":
            c = document.add_paragraph('Diogo Santos-Martins, Leonardo Solis-Vasquez, Andreas F Tillack, Michel F Sanner, Andreas Koch, and Stefano Forli (2021) Accelerating AutoDock4 with GPUs and Gradient-Based Local Search.', style='Heading Four')
            pformat = c.paragraph_format
            pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            c.add_run(' Journal of Chemical Theory and Computation').italic = True
            c.add_run(', ')
            c.add_run('61(8)').bold = True
            c.add_run(': 3891-3898.')
            pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 

        c = document.add_paragraph('Shityakov, S., & Förster, C. (2014). In silico structure-based screening of versatile P-glycoprotein inhibitors using polynomial empirical scoring functions.', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run(' Advances and Applications in Bioinformatics and Chemistry').italic = True
        c.add_run(', ')
        c.add_run('7: ').bold = True
        c.add_run('1.')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('Hopkins, A. L., Groom, C. R., & Alex, A. (2004). Ligand efficiency: a useful metric for lead selection.', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run('  Drug discovery today').italic = True
        c.add_run(', ')
        c.add_run('9(10)').bold = True
        c.add_run(': 430.')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('Bickerton, G. R., Paolini, G. V., Besnard, J., Muresan, S., & Hopkins, A. L. (2012). Quantifying the chemical beauty of drugs.', style='Heading Four')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c.add_run('  Nature chemistry').italic = True
        c.add_run(', ')
        c.add_run('4(2)').bold = True
        c.add_run(': 90-98.')
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('______________________________')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        c = document.add_paragraph('', style = 'Heading Three')
        pformat = c.paragraph_format
        pformat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        j = c.add_run('Attention user: ')
        j.bold = True
        j.italic = True
        c.add_run(' This is an auto-generated writeup for your job, please crosscheck this with the generated output.').italic = True
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        style_five = styles.add_style('Heading Five', WD_STYLE_TYPE.PARAGRAPH)
        style_five.base_style = styles['Heading Five']
        font = style_five.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = True
        font.italic = False
        font.color.rgb = RGBColor(0x1, 0x1, 0x1)
        paragraph = document.add_paragraph('-Team IDv2',style = 'Heading Five')
        pformat = paragraph.paragraph_format
        pformat.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pformat.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.save(os.path.join(results_dir, 'IDv2 Result Summary.docx'))
        return document    
    
    def pdbqt_to_qed(self, pdbqt_path):
        try:
            proc = QProcess()
            args = ["-ipdbqt", pdbqt_path, "-osmi", "-xn"]
            cd = "obabel"
            proc.start(cd, args)
            proc.waitForFinished(-1)

            stdout = proc.readAllStandardOutput().data().decode().strip()
            stderr = proc.readAllStandardError().data().decode().strip()

            # obabel outputs lines like: "SMILES  name"
            lines = stdout.strip().splitlines()
            # take only the very first line

            first_line = lines[0]
            smiles = first_line.split()[0]

            # Build RDKit Mol and compute QED

            mol = Chem.MolFromSmiles(smiles, sanitize=True)
            mol = Chem.AddHs(mol)
            try:
                Chem.SanitizeMol(mol)
            except Chem.SanitizeException:
                pass    
            resultQED = QED.qed(mol)
        except:
            resultQED = 0.0

        return resultQED

    def create_affinity_file_from_pdbqt(self, output_dir, pdbqt_files):
        results = []
        for pdbqt_file in pdbqt_files:
            # Initialize entry with ligand name
            ligand_name = os.path.splitext(os.path.basename(pdbqt_file))[0].replace('_out', '')
            ligand_entry = {'Name': ligand_name}

            # Read PDBQT content
            with open(pdbqt_file, 'r') as f:
                content = f.readlines()

            # Extract torsional energy
            torsional_energy = 'NA'
            try:
                tors_line = next(line for line in content if "TORSDOF" in line)
                tors_value = float(tors_line.split()[1])
                torsional_energy = round(tors_value * 0.3113, 3)
            except (StopIteration, IndexError, ValueError):
                pass

            # Extract binding energy
            energy_value = None
            try:
                energy_line = next(line for line in content if line.startswith('REMARK') and 'Binding energy' in line)
                energy_value = float(energy_line.split('=')[1].strip().split()[0])
            except (StopIteration, IndexError, ValueError):
                try:
                    energy_line = next(line for line in content if line.startswith('REMARK VINA RESULT'))
                    parts = energy_line.strip().split()
                    energy_value = float(parts[3])
                except (StopIteration, IndexError, ValueError):
                    if self.toastSignal:
                        self.toastSignal.emit([2000, f"Skipping {pdbqt_file}: No valid energy line found"])
                    continue

            # Count heavy (non-H) atoms
            nh_count = sum(1 for line in content
                           if line.startswith(('ATOM', 'HETATM')) and ' H ' not in line)

            # Populate energy-derived metrics
            ligand_entry.update({
                'Binding Free Energy (kcal/mol)': energy_value,
                'Ligand Efficiency (kcal/mol/non-H atom)': round(-energy_value / nh_count, 3) if nh_count else 0,
                'pKi': round(-math.log10(math.exp(energy_value * 1000 / (1.986 * 298.15))), 3),
                'Torsional Energy (kcal/mol)': torsional_energy
            })

            # Build RDKit molecule from PDB block and compute QED
            ligand_entry['QED'] = self.pdbqt_to_qed(pdbqt_file)
            results.append(ligand_entry)

        # Write results to CSV
        csv_path = os.path.join(str(output_dir), "IDv2 Results.csv")
        fieldnames = [
            'Name',
            'Binding Free Energy (kcal/mol)',
            'Ligand Efficiency (kcal/mol/non-H atom)',
            'pKi',
            'Torsional Energy (kcal/mol)',
            'QED'
        ]
        with open(csv_path, 'w', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for row in results:
                # Round any floats to 3 decimals
                writer.writerow({k: (round(v, 3) if isinstance(v, float) else v) for k, v in row.items()})

        return csv_path


    def findLigs(self):
        """Return list of ligand files matching the pattern *_out.pdbqt."""
        return glob('*_out.pdbqt')

    def move_ligand_files(self):
        """
        Moves all *_out.pdbqt and (if self.logging is True) *_log.log files from the current working directory
        to a folder structure under the Results folder.
        Returns the path to the Docked_ligands folder.
        """
        results_dir = self.wd + "/Results"
        docked_dir = results_dir + "/Docked Ligands"
        if not os.path.exists(str(docked_dir)):
            os.mkdir(str(docked_dir))
        if self.logging:
            docking_log_dir = docked_dir + "/Docking Log"
            if not os.path.exists(str(docking_log_dir)):
                os.mkdir(str(docking_log_dir))
        for pattern in ["*_out.pdbqt", "*_log.log"]:
            for file in glob(pattern):
                if file.endswith("_log.log"):
                    if self.logging:
                        shutil.move(file, str(docking_log_dir) + os.sep + os.path.basename(file))
                elif file.endswith("_out.pdbqt"):
                    shutil.move(file, str(docked_dir) + os.sep + os.path.basename(file))
        return docked_dir

    def create_top_hits_folder(self, ligand_info_list, top_hits):
        """
        Creates the top hits folder in the Results/Top hits(With Poses) folder by selecting the best compounds,
        copying their pdbqt files (from Docked_ligands) into that folder, splitting each pdbqt file into individual
        conformer files (if multiple conformers exist), and creating an affinity CSV file only for the main pdbqt file.
        Returns a tuple containing the path to the top hits folder and the top hits affinity CSV file.
        """
        if len(ligand_info_list) < 6:
            results_dir = self.wd + "/Results"
            split_dir = results_dir + "/Split Files"
            if not os.path.exists(str(split_dir)):
                os.mkdir(str(split_dir))
            
            sorted_ligs = sorted(ligand_info_list, key=itemgetter('Binding Free Energy (kcal/mol)'))
            top_ligs = sorted_ligs[:int(top_hits)]
            docked_dir = results_dir + "/Docked Ligands"
            
            for lig in top_ligs:
                src = docked_dir + f"/{lig['Name']}_out.pdbqt"
                if os.path.exists(src):
                    dst = str(split_dir) + os.sep + os.path.basename(src)
                    shutil.copy(src, dst)
                    with open(dst, 'r') as f:
                        content = f.read()
                    models = [m.strip() for m in content.split("ENDMDL") if m.strip() != ""]
                    if len(models) > 1:
                        base = os.path.splitext(os.path.basename(src))[0]
                        os.remove(dst)
                        for i, model in enumerate(models, 1):
                            conf_file = os.path.join(str(split_dir), f"{base}_ligand_{i}.pdbqt")
                            with open(conf_file, 'w') as cf:
                                cf.write(model + "\nENDMDL\n")
            return split_dir, None

        else:
            results_dir = self.wd + "/Results"
            split_dir = results_dir + "/Top hits (With Poses)"
            if not os.path.exists(str(split_dir)):
                os.mkdir(str(split_dir))
            
            sorted_ligs = sorted(ligand_info_list, key=itemgetter('Binding Free Energy (kcal/mol)'))
            top_ligs = sorted_ligs[:int(top_hits)]
            docked_dir = results_dir + "/Docked Ligands"
            
            for lig in top_ligs:
                src = docked_dir + f"/{lig['Name']}_out.pdbqt"
                if os.path.exists(src):
                    dst = str(split_dir) + os.sep + os.path.basename(src)
                    shutil.copy(src, dst)
                    with open(dst, 'r') as f:
                        content = f.read()
                    models = [m.strip() for m in content.split("ENDMDL") if m.strip() != ""]
                    if len(models) > 1:
                        base = os.path.splitext(os.path.basename(src))[0]
                        os.remove(dst)
                        for i, model in enumerate(models, 1):
                            conf_file = os.path.join(str(split_dir), f"{base}_ligand_{i}.pdbqt")
                            with open(conf_file, 'w') as cf:
                                cf.write(model + "\nENDMDL\n")
                        main_file = os.path.join(str(split_dir), f"{base}_ligand_1.pdbqt")
                    else:
                        main_file = dst
                    lig['main_file'] = main_file

            top_csv_path = str(split_dir) + os.sep + "IDv2 Results (Top hits).csv"
            with open(top_csv_path, 'w', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(["Name of the ligand", "Binding Free Energy (kcal/mol)", "pKi",
                                "Ligand Efficiency (kcal/mol/non-H atom)", "Torsional Energy (kcal/mol)", "QED"])
                for lig in top_ligs:
                    main_file = lig.get('main_file', "")
                    if not main_file or not os.path.exists(main_file):
                        continue
                    with open(main_file, 'r') as j:
                        ha = j.read()
                    ha_main = ha.split("ENDMDL", 1)[0]
                    ha_lines = ha_main.splitlines()
                    try:
                        energy_line = ha_lines[1]
                        ter = float(energy_line.split(':')[1].split()[0])
                    except Exception:
                        ter = lig.get("Binding Free Energy (kcal/mol)", 0)
                    t = lig.get("Torsional Energy (kcal/mol)", "NA")
                    ki = math.exp(ter * 1000 / (1.986 * 298.15))
                    pki = -math.log(ki, 10)
                    pki = round(pki, 2)
                    nhcount = 0
                    for line in ha_lines:
                        if line.startswith('ATOM') or line.startswith('HETATM'):
                            if ' H ' not in line:
                                nhcount += 1
                        else:
                            nhcount -= 1
                    le = -(ter / nhcount) if nhcount != 0 else 0
                    le = round(le, 4)
                    ligand_name = re.sub(r'\_out\.pdbqt$', '', os.path.basename(main_file))
                    writer.writerow([ligand_name, ter, pki, le, t])
            return split_dir, top_csv_path

    def parse_pdbqt(file_path):
        """
        Parses a PDBQT file and returns an array of atomic positions.
        
        Assumes that atom coordinate lines start with 'ATOM' or 'HETATM' and that
        the x, y, and z coordinates reside in columns 31-38, 39-46, and 47-54 respectively.
        """
        positions = []
        try:
            with open(file_path, "r") as f:
                for line in f:
                    if line.startswith("ATOM") or line.startswith("HETATM"):
                        try:
                            # Columns are based on typical PDB format indices (1-indexed in specification).
                            # In Python (0-indexed), these become:
                            # x: line[30:38], y: line[38:46], z: line[46:54]
                            x = float(line[30:38].strip())
                            y = float(line[38:46].strip())
                            z = float(line[46:54].strip())
                            positions.append([x, y, z])
                        except Exception:
                            pass
        except Exception:
            pass
        
        if not positions:
            return
        
        return np.array(positions)

    def create_rmsd_plot_for_ligand(self, ligand_name, poses_directory, output_path, dpi):
        """
        Finds all pose files for a given ligand in the specified directory (files matching the pattern:
        "<ligand_name>_out_ligand*.pdbqt"), computes the RMSD values of each pose relative to the first pose,
        and plots a bar chart.
        
        This version circumvents MDAnalysis by parsing the PDBQT files directly.
        """
        # Build a search pattern for all pose files of the ligand.
        pattern = os.path.join(poses_directory, f"{ligand_name}_out_ligand*.pdbqt")
        pose_files = glob(pattern)
        
        if not pose_files:
            return

        # Sort pose files by the numeric value following '_conf' (adjust regex if needed)
        def extract_conf_number(filename):
            m = re.search(r'_conf(\d+)', filename)
            return int(m.group(1)) if m else 0
        pose_files.sort(key=extract_conf_number)

        # Load the poses by parsing the PDBQT files
        poses = []
        for pfile in pose_files:
            try:
                positions = continueSingleID.parse_pdbqt(pfile)
                poses.append(positions)
            except Exception:
                pass

        if len(poses) < 2:
            return

        # Compute RMSD values relative to the first pose.
        # It is assumed that every parsed pose has the same number and order of atoms.
        ref = poses[0]
        rmsd_values = []
        for pos in poses:
            if pos.shape != ref.shape:
                continue
            diff = pos - ref
            rmsd = np.sqrt((diff**2).sum() / pos.shape[0])
            rmsd_values.append(rmsd)

        if not rmsd_values:
            return

        # Plot the RMSD values as a bar chart.
        plt.figure(figsize=(8, 6), dpi=dpi)
        indices = list(range(1, len(rmsd_values) + 1))
        plt.bar(indices, rmsd_values, color="skyblue")
        plt.xlabel("Pose Index")
        plt.ylabel("RMSD (Å)")
        plt.title(f"RMSD for ligand {ligand_name} (relative to first pose)")
        plt.tight_layout()
        plt.savefig(output_path, dpi=dpi)
        plt.close()

    def create_binding_energy_plot(self, energies, output_path, dpi):
        """
        Create a histogram with KDE showing the distribution of binding energies.
        """
        plt.figure(figsize=(8, 6), dpi=self.dpi)
        sns.histplot(energies, kde=True, bins=15, color="skyblue")
        plt.title("Distribution of Binding Energies")
        plt.xlabel("Binding Free Energy (kcal/mol)")
        plt.ylabel("Frequency")
        plt.tight_layout()
        plt.savefig(output_path, dpi=dpi)
        plt.close()

    def run(self):
        """
        Main method executed in this thread.
        Steps:
          1. Find docking result ligand files.
          2. Create the overall affinity CSV file (saved in Results folder).
          3. Parse the CSV to obtain ligand info.
          3.5 If >6 compounds, prompt the user for the number of top hits; otherwise, use all.
          4. Adjust mode if >1000 compounds.
          5. Move ligand files to Results/Docked ligands (and _log.log files to Docked ligands/Docking log if self.logging is True) and move target files (IF GPU based docking) to Target map files.
          6. Create the docking document in Results.
          7. Create plots in Results.
          8. Create the top hits folder in Results/Split_files (copy top hits files and create top hits CSV).
        """
        results_dir = self.wd + "/Results"
        if not os.path.exists(str(results_dir)):
            os.mkdir(str(results_dir))
        pdbqt_files = self.findLigs()
        if not pdbqt_files:
            self.finishedSignal.emit(["❌ No docking result files found."])
            return
        receptor_name = self.findTarget()
        self.progressBarSignal.emit(75)
        self.progressLabelSignal.emit("<b><center>Progress: 75% (Creating overall affinity results file...)</center></b>")
        try:
            affinity_csv = self.create_affinity_file_from_pdbqt(results_dir, pdbqt_files)
        except Exception:           
            self.finishedSignal.emit([f"❌ Error creating affinity file."])
            return

        self.progressBarSignal.emit(80)
        self.progressLabelSignal.emit("<b><center>Progress: 80% (Parsing overall affinity results...)</center></b>")
        ligand_info_list = []
        try:
            with open(affinity_csv, 'r', newline='') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    row['Binding Free Energy (kcal/mol)'] = float(row['Binding Free Energy (kcal/mol)'])
                    row['pKi'] = float(row['pKi'])
                    row['Ligand Efficiency (kcal/mol/non-H atom)'] = float(row['Ligand Efficiency (kcal/mol/non-H atom)'])
                    if row['Torsional Energy (kcal/mol)'].strip().upper() != 'NA':
                        row['Torsional Energy (kcal/mol)'] = float(row['Torsional Energy (kcal/mol)'])
                    else:
                        row['Torsional Energy (kcal/mol)'] = 'NA'
                    if row['QED'].strip().upper() != 'NA':
                        row['QED'] = float(row['QED'])
                    else:
                        row['QED'] = 'NA'

                    ligand_info_list.append(row)

        except Exception:
            self.finishedSignal.emit([f"❌ Error reading affinity file"])
            return
        
        if len(ligand_info_list) > 10:
            self.requestInput.emit(len(ligand_info_list))
            while not self.inputReceived:
                QThread.msleep(100)  # sleep 100 ms
                continue

            self.top_hits = self.top_hits
            self.toastSignal.emit([2000, f"Top hits set to: {self.top_hits}"])

        else:
            self.top_hits = len(ligand_info_list)

        if len(ligand_info_list) > 1000:
            self.mode = "virtual_screening"
        else:
            self.mode = "docking"

        self.progressBarSignal.emit(85)
        self.progressLabelSignal.emit("<b><center>Progress: 85% (Moving docked ligand files...)</center></b>")
        docked_dir = self.move_ligand_files()
        self.toastSignal.emit([2000, f"Docked ligand files moved to {docked_dir}"])
        if self.method == "AutoDock-GPU [2]":
            receptor = receptor_name[0].split('.')[0]
            self.moveTarget(receptor)

        self.progressBarSignal.emit(90)
        self.progressLabelSignal.emit("<b><center>Progress: 90% (Creating docking results document...)</center></b>")
        try:
            doc = self.create_docking_document(receptor_name, ligand_info_list, results_dir, self.mode, self.top_hits, self.method)
        except Exception:
            self.finishedSignal.emit([f"❌ Error creating docking document"])
            return
        self.toastSignal.emit([2000, "Docking results document created."])

        self.progressBarSignal.emit(95)
        self.progressLabelSignal.emit("<b><center>Progress: 95% (Creating top hits folder...)</center></b>")
        split_folder, top_csv_path = self.create_top_hits_folder(ligand_info_list, self.top_hits)
        if len(ligand_info_list) > 10:
            self.toastSignal.emit([2000, f"Top hits files and affinity file created at {split_folder}"])

        self.progressBarSignal.emit(99)
        self.progressLabelSignal.emit("<b><center>Progress: 99% (Creating plots...)</center></b>")
        plots_dir = results_dir + "/Plots"
        if not os.path.exists(str(plots_dir)):
            os.mkdir(str(plots_dir))
        if len(ligand_info_list) <= 10:
            # Extract the ligand name by removing '_out' from the file name.
            for ligand in ligand_info_list:
                ligand_name = ligand['Name']
                rmsd_plot_path = str(plots_dir) + os.sep + f"{ligand_name}_RMSD_Heatmap.png"
                # Use the docked directory as the location where pose files are stored.
                dpi = self.dpi
                self.create_rmsd_plot_for_ligand(ligand_name, split_folder, rmsd_plot_path, dpi)
            self.toastSignal.emit([2000, f"RMSD plot created at {os.path.dirname(rmsd_plot_path)}"])
        else:
            energies = [lig['Binding Free Energy (kcal/mol)'] for lig in ligand_info_list]
            be_plot_path = str(plots_dir) + os.sep + "Binding_Energy_Distribution.png"
            self.create_binding_energy_plot(energies, be_plot_path, dpi=self.dpi)
            self.toastSignal.emit([2000, f"Binding energy distribution plot created at {be_plot_path}"])

        self.toastSignal.emit([2000, "Docking processing completed successfully!"])
        
        self.finishedSignal.emit(["DONE", self.mode])
